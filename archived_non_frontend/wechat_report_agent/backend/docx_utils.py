import os
import tempfile
import logging

logger = logging.getLogger('backend.docx')


def simple_generate_docx(content_dict, out_path):
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import Pt
    except Exception as e:
        logger.debug('python-docx not available: %s', e)
        raise
    doc = Document()

    try:
        normal = doc.styles['Normal']
        font = normal.font
        font.name = 'Times New Roman'
        rpr = normal.element.rPr
        if rpr is None:
            from docx.oxml import OxmlElement
            rpr = OxmlElement('w:rPr')
            normal.element.append(rpr)
        normal.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        font.size = Pt(12)
    except Exception:
        logger.debug('failed to set Normal style fonts')

    def _clean_text_for_docx(text: str) -> str:
        try:
            import re
            if text is None:
                return ''
            s = str(text)
            s = re.sub(r'\*{2,}', '', s)
            s = re.sub(r'_{2,}', '', s)
            s = re.sub(r'`+', '', s)
            s = re.sub(r'^[\-=*_]{3,}$', '', s, flags=re.M)
            s = re.sub(r'(?:\n[\-\*]{2,}\n)+', '\n', s)
            s = re.sub(r'[ \t\u00A0]{2,}', ' ', s)
            def _fw2hw(m):
                return ''.join(chr(ord(c) - 0xFF10 + ord('0')) if '\uFF10' <= c <= '\uFF19' else c for c in m.group(0))
            s = re.sub(r'[\uFF10-\uFF19]+', _fw2hw, s)
            s = re.sub(r'^\s*好的[，,\s\S]{0,120}?(?:生成的(?:新闻)?片段|生成的(?:内容|文本|片段|结果))[:：\-\s]*', '', s, flags=re.I|re.M)
            s = re.sub(r'^\s*好的[，,\s\S]{0,120}?：\s*---\s*', '', s, flags=re.I|re.M)
            s = re.sub(r'(?m)^[\-\*]{2,}\s*$', '\n', s)
            s = re.sub(r'\n{2,}', '\n\n', s)
            s = s.replace('：', ':').replace('（', '(').replace('）', ')')
            s = s.replace('--', '—').replace('—–', '—')
            s = s.replace('\u2013', '—').replace('\u2014', '—')
            s = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]+', '', s)
            s = s.strip()
            return s
        except Exception:
            return str(text or '')

    def _split_cn_en_runs(text: str):
        import re
        if not text:
            return []
        cn_re = re.compile(r'[\u4e00-\u9fff\u3000-\u303f\uff00-\uffef]+')
        parts = []
        idx = 0
        for m in cn_re.finditer(text):
            if m.start() > idx:
                parts.append((text[idx:m.start()], False))
            parts.append((m.group(0), True))
            idx = m.end()
        if idx < len(text):
            parts.append((text[idx:], False))
        return [(p, is_cn) for p, is_cn in parts if p]

    def _add_heading(text, level=1):
        p = doc.add_heading(level=level)
        run = p.add_run(_clean_text_for_docx(text))
        try:
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            run.font.size = Pt(16 if level == 1 else 14)
        except Exception:
            pass

    def _add_paragraph(text):
        p = doc.add_paragraph()
        s = _clean_text_for_docx(text)
        for seg, is_cn in _split_cn_en_runs(s):
            r = p.add_run(seg)
            try:
                if is_cn:
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    r.font.name = 'Times New Roman'
                else:
                    r.font.name = 'Times New Roman'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                r.font.size = Pt(12)
            except Exception:
                pass

    _add_heading('AI 生成的报告（简易版）', level=1)
    for k, items in content_dict.items():
        _add_heading(k, level=2)
        if not items:
            _add_paragraph('(无内容)')
        else:
            for it in items:
                if isinstance(it, dict):
                    title = it.get('title', '')
                    summary = it.get('summary', '')
                    if title:
                        p = doc.add_paragraph()
                        for seg, is_cn in _split_cn_en_runs(_clean_text_for_docx(title) + '\n'):
                            r1 = p.add_run(seg)
                            try:
                                r1.font.bold = True
                                if is_cn:
                                    r1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                                else:
                                    r1.font.name = 'Times New Roman'
                                r1.font.size = Pt(12)
                            except Exception:
                                pass
                        for seg, is_cn in _split_cn_en_runs(_clean_text_for_docx(summary)):
                            r2 = p.add_run(seg)
                            try:
                                if is_cn:
                                    r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                                else:
                                    r2.font.name = 'Times New Roman'
                                r2.font.size = Pt(12)
                            except Exception:
                                pass
                    else:
                        _add_paragraph(summary)
                else:
                    _add_paragraph(it)

    doc.save(out_path)
