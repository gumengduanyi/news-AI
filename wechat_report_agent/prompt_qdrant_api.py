"""本文件为本地测试提供一个最小化的 Flask 后端。

提供的接口：
- GET  /api/validate-env
- POST /api/test-models
- GET/POST /api/generate-report-debug

该文件尽量依赖最少的第三方库，如果内部模块不可用会回退到本地的模拟实现（mock）。
"""
import os
import json
import datetime
import tempfile
import time
from flask import Flask, request, jsonify, send_file
import logging

import sys
import sqlite3
from typing import List, Dict, Any

# 可选导入（用于 RAG 功能），若不可用会回退为 None
try:
    from qdrant_client import QdrantClient
    from qdrant_client.models import ScoredPoint
except Exception:
    QdrantClient = None

try:
    from sentence_transformers import SentenceTransformer
except Exception:
    SentenceTransformer = None
# NOTE: defer importing ai provider bindings (e.g. call_deepseek_embed) until after
# we ensure instance/config.json has been loaded into os.environ (see loader below).

# 使用模块局部路径设置以确保本地导入正常，避免遗留代码查找问题
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
if BASE_DIR not in sys.path:
    sys.path.insert(0, BASE_DIR)
PARENT = os.path.abspath(os.path.join(BASE_DIR, '..'))
if PARENT not in sys.path:
    sys.path.insert(0, PARENT)

# --- load local instance config (optional) ---------------------------------
# If there's an instance/config.json file, load keys from it into os.environ
# This allows developers to store API keys locally without exporting them every time.
try:
    _config_path = os.path.join(BASE_DIR, 'instance', 'config.json')
    if os.path.exists(_config_path):
        with open(_config_path, 'r', encoding='utf-8') as _cf:
            _cfg = json.load(_cf)
            for _k, _v in _cfg.items():
                # only set if not already in environment
                if isinstance(_v, str) and not os.environ.get(_k):
                    os.environ[_k] = _v
        # we can't use logger here because it may not be initialized yet; rely on later log
except Exception:
    pass

try:
    from wechat_report_agent.ai_api import call_ai
except Exception as e:
    # will use a local fallback if ai_api not available
    call_ai = None
    # defer logging setup until logger exists
    _import_callai_err = str(e)

try:
    from wechat_report_agent.src.render_word_report import generate_ai_report, render_report
except Exception as e:
    generate_ai_report = None
    # render_report is optional (used when prompt_template.content points to a .docx template file)
    try:
        render_report = None
    except Exception:
        render_report = None
    _import_generate_err = str(e)

app = Flask(__name__)

# --- 简易 CORS 支持 （允许本地前端在不同端口访问） -----------------
# This minimal CORS handler will add the necessary headers to responses
# and respond to OPTIONS preflight requests. It's intentionally simple so
# the project doesn't require additional dependencies in dev.
@app.after_request
def _add_cors_headers(response):
    # Allow requests from the typical local dev UI (adjust as needed)
    allowed_origin = os.environ.get('ALLOW_ORIGIN', '*')
    response.headers['Access-Control-Allow-Origin'] = allowed_origin
    response.headers['Access-Control-Allow-Credentials'] = 'true'
    response.headers['Access-Control-Allow-Headers'] = 'Authorization,Content-Type,Accept,Origin,User-Agent'
    response.headers['Access-Control-Allow-Methods'] = 'GET,POST,PUT,DELETE,OPTIONS'
    # expose ai counts header used by frontend
    response.headers['Access-Control-Expose-Headers'] = 'X-AI-COUNTS'
    return response


@app.route('/', methods=['OPTIONS'])
def _handle_options_root():
    # generic OPTIONS handler for root path
    return ('', 200)


# 为本模块设置结构化日志
logger = logging.getLogger('prompt_qdrant_api')
if not logger.handlers:
    handler = logging.StreamHandler()
    handler.setFormatter(logging.Formatter('%(asctime)s %(levelname)s %(message)s'))
    logger.addHandler(handler)
logger.setLevel(logging.INFO)

# --- load local instance config (optional) ---------------------------------
# If there's an instance/config.json file, load keys from it into os.environ
# This allows developers to store API keys locally without exporting them every time.
try:
    _config_path = os.path.join(BASE_DIR, 'instance', 'config.json')
    if os.path.exists(_config_path):
        with open(_config_path, 'r', encoding='utf-8') as _cf:
            _cfg = json.load(_cf)
            for _k, _v in _cfg.items():
                # only set if not already in environment
                if isinstance(_v, str) and not os.environ.get(_k):
                    os.environ[_k] = _v
        logger.info('Loaded instance config from %s', _config_path)
except Exception as _e:
    logger.warning('Failed to load instance config %s: %s', _config_path, _e)

# 如果导入失败，记录警告信息以便排查
if ' _import_callai_err' in globals():
    logger.warning('无法导入 call_ai: %s', globals().get('_import_callai_err'))
if ' _import_generate_err' in globals():
    logger.warning('无法导入 generate_ai_report: %s', globals().get('_import_generate_err'))


# Local fallback for call_ai when the real provider is not installed/available.
def local_call_ai_fallback(model, prompt):
    """简易的本地回退实现：尝试从提示词中抽取要点或生成小段模拟响应，便于离线测试。

    行为说明：
    - 若提示中包含要求返回严格 JSON 的字样，返回一个空的五键 JSON 模板。
    - 尝试从提示文本中提取看起来像要点的行（数字序号、破折号或短句），作为响应返回。
    - 否则返回默认的模拟文本。
    """
    logger.info('local_call_ai_fallback invoked for model=%s', model)
    # 如果提示包含严格返回 JSON 的要求，则尝试从提示或材料中抽取要点并分配到五个键中
    if '请严格返回JSON' in (prompt or '') or '严格转换为可解析的JSON' in (prompt or ''):
        expected_keys = ["core_news", "技术前沿", "产业动态", "政策法规", "应用实例"]
        import re
        text = prompt or ''
        m = re.search(r'(?:输入文本：|材料：)([\s\S]*)', text)
        body = (m.group(1) if m else text).strip()

        # 切分候选要点（按换行、句号或分号），去掉空行和短碎片
        parts = []
        for part in re.split(r'[\n。；;]+', body):
            s = part.strip()
            if not s:
                continue
            s = re.sub(r'^[\-\*\d\.\s]+', '', s)
            if len(s) >= 6:
                parts.append(s)

        if not parts:
            parts = [
                '示例：AI 行业发布重大模型更新',
                '示例：企业宣布新一轮融资',
                '示例：监管机构发布相关指引',
                '示例：行业出现新型应用场景',
                '示例：研究机构发布重要论文'
            ]

        schema = {k: [] for k in expected_keys}
        i = 0
        for p in parts:
            key = expected_keys[i % len(expected_keys)]
            if len(schema[key]) < 3:
                schema[key].append(p)
            i += 1
        return json.dumps(schema, ensure_ascii=False)

    # 尝试抽取看起来像要点的行
    lines = []
    for ln in (prompt or '').splitlines():
        s = ln.strip()
        if not s:
            continue
        if s[0].isdigit() or s.startswith('-') or s.startswith('*'):
            lines.append(s.lstrip('-*0123456789. ').strip())
        elif len(s) < 120:
            lines.append(s)
    if lines:
        return '\n'.join(lines[:8])

    # 默认短响应
    return '这是本地模拟的模型响应。'

# 如果未导入真实的 call_ai，则指向本地回退实现，避免在调用处做分支判断
if call_ai is None:
    # only auto-install local fallback if explicitly enabled in env
    if os.environ.get('DEV_CALLAI_FALLBACK') == '1':
        logger.warning('call_ai not available and DEV_CALLAI_FALLBACK=1: installing local_call_ai_fallback')
        call_ai = local_call_ai_fallback
    else:
        logger.warning('call_ai not available and DEV_CALLAI_FALLBACK not set: local fallback disabled')
    # instance config already loaded at module top; nothing to do here


def require_token(func):
    """鉴权装饰器：检查 REPORT_API_TOKEN（env）或请求头/查询参数/cookie 中的 token。

    如果在本机开发环境设置了 DEV_AUTH_DISABLED=1，则允许来自本机（127.0.0.1/::1/localhost）的请求跳过鉴权。
    """
    from functools import wraps

    @wraps(func)
    def wrapper(*args, **kwargs):
        # 本地开发绕过
        if os.environ.get('DEV_AUTH_DISABLED') == '1':
            remote = request.remote_addr or ''
            if remote.startswith('127.') or remote == '::1' or remote == 'localhost':
                logger.warning('DEV_AUTH_DISABLED=1: skipping auth checks for local request from %s', remote)
                return func(*args, **kwargs)
            else:
                logger.info('DEV_AUTH_DISABLED=1 present but request remote_addr=%s not local; enforcing auth', remote)

        expected = os.environ.get('REPORT_API_TOKEN')
        if not expected:
            # 强制要求服务器配置 REPORT_API_TOKEN
            logger.error('REQUEST REJECTED: server missing REPORT_API_TOKEN (server misconfigured)')
            return jsonify({'status': 'error', 'error': 'server_misconfigured_missing_token'}), 500

        # 检查 Authorization 头部
        auth = request.headers.get('Authorization', '')
        token = None
        if auth.startswith('Bearer '):
            token = auth.split(' ', 1)[1].strip()

        # 回退到查询参数 token
        if not token:
            token = request.args.get('token')
            # 如果仍无 token，兼容从 cookie 中读取（前端可能在同源请求时把 token 存在 cookie 中）
            if not token:
                try:
                    token = request.cookies.get('TOKEN')
                except Exception:
                    token = None

        if token != expected:
            # 记录日志但不要泄露完整 token（仅长度/掩码）
            provided = token or '<none>'
            masked = (provided[:4] + '...' + provided[-4:]) if len(provided) > 8 else provided
            logger.warning('AUTH FAILED: provided token=%s expected=***', masked)
            return jsonify({'status': 'error', 'error': 'invalid_or_missing_token'}), 401
        return func(*args, **kwargs)

    return wrapper


def _write_instance_config(updates: dict):
    """Write selected keys to instance/config.json and set file perms to 600.

    Only string values are written. Existing keys will be updated/added.
    """
    inst_dir = os.path.join(BASE_DIR, 'instance')
    os.makedirs(inst_dir, exist_ok=True)
    cfg_path = os.path.join(inst_dir, 'config.json')
    try:
        if os.path.exists(cfg_path):
            with open(cfg_path, 'r', encoding='utf-8') as f:
                cur = json.load(f)
        else:
            cur = {}
    except Exception:
        cur = {}
    # merge
    for k, v in updates.items():
        if isinstance(v, str):
            cur[k] = v
    # write atomically
    tmp = cfg_path + '.tmp'
    with open(tmp, 'w', encoding='utf-8') as f:
        json.dump(cur, f, ensure_ascii=False, indent=2)
    os.replace(tmp, cfg_path)
    try:
        os.chmod(cfg_path, 0o600)
    except Exception:
        logger.debug('failed to chmod %s', cfg_path)
    # inject into env
    for k, v in cur.items():
        if isinstance(v, str):
            os.environ[k] = v


@app.route('/api/set-keys', methods=['POST'])
@require_token
def api_set_keys():
    """受保护的运行时接口：写入 instance/config.json 并注入到环境变量。

    请求 body 应为 JSON，包含要写入的键值对，如 {"DEEPSEEK_API_KEY":"xxx","REPORT_API_TOKEN":"yyy"}
    仅写入字符串值，且推荐仅在本地开发环境使用。
    """
    data = request.json or {}
    if not isinstance(data, dict):
        return jsonify({'status': 'error', 'error': 'invalid_payload'}), 400

    # whitelist keys to avoid accidental secrets being stored
    allowed_prefixes = ('DEEPSEEK', 'DOUBAO', 'ZHIPUAI', 'REPORT', 'QDRANT', 'DEEPSEEK_EMBED')
    updates = {k: v for k, v in data.items() if isinstance(k, str) and isinstance(v, str) and any(k.startswith(p) for p in allowed_prefixes)}
    if not updates:
        return jsonify({'status': 'error', 'error': 'no_valid_keys_provided'}), 400

    try:
        _write_instance_config(updates)
        return jsonify({'status': 'ok', 'written': list(updates.keys())})
    except Exception as e:
        logger.exception('failed to write instance config')
        return jsonify({'status': 'error', 'error': str(e)}), 500


# Module-level cache to hold last AI outputs for debugging (kept in memory only)
_last_ai_debug = {'ai_raw': None, 'ai_content': None, 'model': None}


# 全局异常处理器：记录未捕获异常并返回一致的 JSON 错误响应。
# 通过环境变量 PROMPT_QDRANT_DEBUG_ERRORS=1 可在响应中包含 traceback（仅用于本地调试），生产环境请勿启用。
import traceback as _traceback
@app.errorhandler(Exception)
def _handle_unhandled_exception(e):
    logger.exception('Unhandled exception in request: %s', e)
    if os.environ.get('PROMPT_QDRANT_DEBUG_ERRORS') == '1':
        tb = _traceback.format_exc()
        return jsonify({'status': 'error', 'error': str(e), 'traceback': tb}), 500
    return jsonify({'status': 'error', 'error': 'internal_server_error'}), 500


@app.route('/api/reload-instance-config', methods=['POST', 'GET'])
@require_token
def api_reload_instance_config():
    """热加载 instance/config.json：重新读取文件并注入到 os.environ。

    仅在本地开发/受保护环境使用。返回已注入的键名和掩码后的示例值（仅用于确认注入是否成功）。
    """
    inst_path = os.path.join(BASE_DIR, 'instance', 'config.json')
    if not os.path.exists(inst_path):
        return jsonify({'status': 'error', 'error': 'config_not_found', 'path': inst_path}), 404
    try:
        with open(inst_path, 'r', encoding='utf-8') as f:
            cfg = json.load(f)
    except Exception as e:
        logger.exception('failed to read instance config')
        return jsonify({'status': 'error', 'error': 'read_failed', 'details': str(e)}), 500

    injected = {}
    for k, v in cfg.items():
        if isinstance(k, str) and isinstance(v, str):
            os.environ[k] = v
            # mask value for response (show first 4 + last 4 chars if long)
            if len(v) > 12:
                masked = v[:4] + '...' + v[-4:]
            else:
                masked = v[:2] + '...' if len(v) > 4 else '***'
            injected[k] = masked

    logger.info('Reloaded instance config and injected keys: %s', ','.join(list(injected.keys())))
    return jsonify({'status': 'ok', 'injected': injected})

# --- 简易 docx 生成器（模块级，便于测试导入） -----------------
def simple_generate_docx(content_dict, out_path):
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import Pt
    except Exception as e:
        logger.debug('python-docx not available: %s', e)
        raise
    doc = Document()

    # Set default Normal style fonts (Latin and East Asia)
    try:
        normal = doc.styles['Normal']
        font = normal.font
        # Latin font
        font.name = 'Times New Roman'
        # East Asia font (for Chinese)
        rpr = normal.element.rPr
        if rpr is None:
            from docx.oxml import OxmlElement
            rpr = OxmlElement('w:rPr')
            normal.element.append(rpr)
        normal.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        font.size = Pt(12)
    except Exception:
        # ignore style setting failures; continue to per-run setting
        logger.debug('failed to set Normal style fonts')

    def _clean_text_for_docx(text: str) -> str:
        """Clean text to reduce markdown artifacts, control chars and odd punctuation that
        may cause display issues in WPS/Word. Returns cleaned unicode string."""
        try:
            import re
            if text is None:
                return ''
            s = str(text)
            # remove common markdown markers like **bold**, __, ``, and repeated asterisks
            s = re.sub(r'\*{2,}', '', s)
            s = re.sub(r'_{2,}', '', s)
            s = re.sub(r'`+', '', s)
            # remove sequences like --- or *** used as separators
            s = re.sub(r'^[\-=*_]{3,}$', '', s, flags=re.M)
            # remove excessive internal markers like ---, *** or ---片段---
            s = re.sub(r'(?:\n[\-\*]{2,}\n)+', '\n', s)
            # replace multiple whitespace with single space but preserve paragraph breaks
            s = re.sub(r'[ \t\u00A0]{2,}', ' ', s)
            # convert fullwidth digits to ASCII digits (０１２３ -> 0123)
            def _fw2hw(m):
                return ''.join(chr(ord(c) - 0xFF10 + ord('0')) if '\uFF10' <= c <= '\uFF19' else c for c in m.group(0))
            s = re.sub(r'[\uFF10-\uFF19]+', _fw2hw, s)
            # remove template intro lines often inserted by prompt templates
            # variants: "好的，这是根据您提供的标题生成的片段：", "好的，这是根据您提供的信息生成的新闻片段：" etc.
            s = re.sub(r'^\s*好的[，,\s\S]{0,120}?(?:生成的(?:新闻)?片段|生成的(?:内容|文本|片段|结果))[:：\-\s]*', '', s, flags=re.I|re.M)
            s = re.sub(r'^\s*好的[，,\s\S]{0,120}?：\s*---\s*', '', s, flags=re.I|re.M)
            # remove leading '---' separators and blocks of dashes/asterisks
            s = re.sub(r'(?m)^[\-\*]{2,}\s*$', '\n', s)
            # remove parenthetical placeholder fragments like (可选材料未提供具体内容...)
            s = re.sub(r'\([^\)]*(可选材料|可直接使用|可选材料|示例|片段|片段:)[^\)]*\)', '', s)
            # drop lines that are very short and look like placeholders (e.g. '---', '片段', '* *')
            out_lines = []
            for ln in s.splitlines():
                lns = ln.strip()
                if not lns:
                    continue
                if len(lns) <= 4 and re.match(r'^[\-*\*\s\u2014\u2013\u2022]+$', lns):
                    continue
                # remove bracketed stage or video/script markers like （视频片段） or 【动感音乐响起，...】
                if re.match(r'^[\(（\[【].{0,200}[\)）\]】]$', lns):
                    # if the bracketed content is short and not a real sentence, drop it
                    if len(lns) < 200 and not re.search(r'[。\.\?\！!\?]', lns):
                        continue
                if lns.startswith('**') or lns.endswith('**'):
                    # remove surrounding asterisks
                    lns = lns.strip('* ').strip()
                    if not lns:
                        continue
                # ignore pure '示例' or '片段' lines
                if re.match(r'^[示例片段示范示例：:]+$', lns):
                    continue
                out_lines.append(lns)
            s = '\n'.join(out_lines)
            # collapse multiple blank lines
            s = re.sub(r'\n{2,}', '\n\n', s)
            # normalize some fullwidth punctuation to ASCII where safe
            s = s.replace('：', ':').replace('（', '(').replace('）', ')')
            # normalize different dash types to em-dash
            s = s.replace('--', '—').replace('—–', '—')
            s = s.replace('\u2013', '—').replace('\u2014', '—')
            # remove C0 control chars except newline and tab
            s = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]+', '', s)
            # strip weird leading/trailing markers
            s = s.strip()
            return s
        except Exception:
            return str(text or '')


    def _split_cn_en_runs(text: str):
        """Split text into segments that are mostly Chinese (CJK) vs non-Chinese.

        Yields tuples (segment, is_chinese_bool).
        Keeps punctuation with the segment it belongs to.
        """
        import re
        if not text:
            return []
        # Chinese (CJK unified) unicode ranges roughly: \u4e00-\u9fff, plus punctuation range
        cn_re = re.compile(r'[\u4e00-\u9fff\u3000-\u303f\uff00-\uffef]+')
        parts = []
        idx = 0
        for m in cn_re.finditer(text):
            # non-cn before
            if m.start() > idx:
                parts.append((text[idx:m.start()], False))
            parts.append((m.group(0), True))
            idx = m.end()
        if idx < len(text):
            parts.append((text[idx:], False))
        # flatten empty
        return [(p, is_cn) for p, is_cn in parts if p]


    def _add_heading(text, level=1):
        p = doc.add_heading(level=level)
        run = p.add_run(_clean_text_for_docx(text))
        try:
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            run.font.size = Pt(16 if level == 1 else 14)
        except Exception:
            pass

    def _add_paragraph(text):
        p = doc.add_paragraph()
        s = _clean_text_for_docx(text)
        # split into Chinese / non-Chinese runs to set appropriate fonts
        for seg, is_cn in _split_cn_en_runs(s):
            r = p.add_run(seg)
            try:
                if is_cn:
                    # Chinese (East Asia) font
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    r.font.name = 'Times New Roman'  # keep Latin name for compatibility
                else:
                    r.font.name = 'Times New Roman'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                r.font.size = Pt(12)
            except Exception:
                pass

    _add_heading('AI 生成的报告（简易版）', level=1)
    for k, items in content_dict.items():
        _add_heading(k, level=2)
        if not items:
            _add_paragraph('(无内容)')
        else:
            for it in items:
                # items may be dicts (with title/summary) or strings
                if isinstance(it, dict):
                    title = it.get('title', '')
                    summary = it.get('summary', '')
                    if title:
                        # title as bold run
                        p = doc.add_paragraph()
                        # title (may contain mixed text) -> split into runs
                        for seg, is_cn in _split_cn_en_runs(_clean_text_for_docx(title) + '\n'):
                            r1 = p.add_run(seg)
                            try:
                                r1.font.bold = True
                                if is_cn:
                                    r1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                                else:
                                    r1.font.name = 'Times New Roman'
                                r1.font.size = Pt(12)
                            except Exception:
                                pass
                        # summary -> split into runs
                        for seg, is_cn in _split_cn_en_runs(_clean_text_for_docx(summary)):
                            r2 = p.add_run(seg)
                            try:
                                if is_cn:
                                    r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                                else:
                                    r2.font.name = 'Times New Roman'
                                r2.font.size = Pt(12)
                            except Exception:
                                pass
                    else:
                        _add_paragraph(summary)
                else:
                    _add_paragraph(it)

    # Save document
    doc.save(out_path)


def normalize_ai_content_for_render(ai_content):
    """Normalize ai_content into a shape expected by renderers.

    - 'core_news' stays as a list of strings.
    - other sections become a list of dicts with 'title' and 'summary'.
    This avoids templates skipping content when AI returns strings or varied dict shapes.
    """
    expected_sections = ["技术前沿", "产业动态", "政策法规", "应用实例"]
    out = {}
    # core_news: ensure list of strings
    core = ai_content.get('core_news') if isinstance(ai_content, dict) else []
    core_out = []
    for it in (core or []):
        if isinstance(it, str):
            core_out.append(it)
        elif isinstance(it, dict):
            # try to extract a readable title
            title = it.get('title') or it.get('name') or it.get('headline') or it.get('summary') or ''
            if title:
                core_out.append(title)
            else:
                core_out.append(json.dumps(it, ensure_ascii=False))
        else:
            core_out.append(str(it))
    out['core_news'] = core_out

    for sec in expected_sections:
        sec_items = ai_content.get(sec) if isinstance(ai_content, dict) else []
        norm = []
        for it in (sec_items or []):
            if isinstance(it, str):
                # Treat plain string items as both title and summary so the DOCX will show content
                item = {'title': it, 'summary': it}
                # compatibility: also provide Chinese field names if frontend expects them
                item['标题'] = it
                item['摘要'] = it
                norm.append(item)
            elif isinstance(it, dict):
                title = it.get('title') or it.get('name') or it.get('headline') or ''
                summary = it.get('summary') or it.get('content') or it.get('description') or ''
                # if title empty but summary present, use a prefix of summary as title
                if not title and summary:
                    title = summary if len(summary) <= 120 else summary[:120]
                item = {'title': title, 'summary': summary}
                # also fill Chinese keys for compatibility with some frontends
                item['标题'] = title
                item['摘要'] = summary
                norm.append(item)
            else:
                norm.append({'title': str(it), 'summary': ''})
        out[sec] = norm

    # post-process summaries: remove common AI-prefixes and strip verbatim sentences from combined_material
    try:
        import re
        cm_text = ai_content.get('combined_material') if isinstance(ai_content, dict) else ''
        cm_sents = []
        if cm_text and isinstance(cm_text, str):
            cm = cm_text
            for sep in ('。', '.', '！', '!', '?', '？', ';', '；', '\n'):
                cm = cm.replace(sep, '。')
            cm_sents = [s.strip() for s in cm.split('。') if s.strip()]

        for sec in expected_sections:
            for item in out.get(sec, []):
                s = item.get('summary') or ''
                if not s:
                    continue
                # remove markers like 'AI 生成：', 'AI Generated:' and '相关片段：' anywhere
                s = re.sub(r'(AI\s*生成：|AI Generated:|相关片段：)\s*', '', s, flags=re.I)
                # remove exact sentences that appear in combined_material
                try:
                    for sent in cm_sents:
                        if sent and sent in s:
                            s = s.replace(sent, '').strip()
                    # simple dedupe by sentences
                    parts = re.split(r'([。\.\!\?\？\n])', s)
                    new_parts = []
                    seen = set()
                    for i in range(0, len(parts), 2):
                        sent = parts[i].strip()
                        sep = parts[i+1] if i+1 < len(parts) else ''
                        full = (sent + sep).strip()
                        if not full:
                            continue
                        if full in seen:
                            continue
                        seen.add(full)
                        new_parts.append(full)
                    s = ' '.join(new_parts).strip()
                except Exception:
                    pass
                # fallback to placeholder if empty
                if not s:
                    s = '（该条暂无详细摘要 — 请检查 AI 输出或后端日志）'
                item['summary'] = s
    except Exception:
        logger.exception('post-process summaries failed')

    logger.info('normalize_ai_content_for_render: counts=%s', {k: len(v) for k, v in out.items()})
    return out


def ensure_structured_ai_response(model, ai_output):
    """确保 ai_output 被规范为一个包含五个预期键且每个键对应列表的字典。

    处理流程：
    - 若 ai_output 已是 dict：补全缺失键并返回。
    - 若 ai_output 为字符串：尝试 json.loads；若失败，尝试提取 JSON 子串解析。
    - 若仍失败且 `call_ai` 可用，则请求模型将文本转换为严格的 JSON。
    - 若所有方法都失败，返回空的 schema（每个键对应空列表）。
    """
    expected_keys = ["core_news", "技术前沿", "产业动态", "政策法规", "应用实例"]

    # 辅助函数：规范化字典，保证五个键都存在且值为列表
    def _normalize(d):
        out = {}
        for k in expected_keys:
            v = d.get(k, []) if isinstance(d, dict) else []
            if v is None:
                v = []
            # ensure list
            if not isinstance(v, list):
                v = [v]
            out[k] = v
        return out

    logger.info('ensure_structured_ai_response called for model=%s, ai_output_type=%s', model, type(ai_output).__name__)
    import re
    def _mask_snippet(s, length=1000):
        try:
            t = s if isinstance(s, str) else json.dumps(s, ensure_ascii=False)
        except Exception:
            t = str(s)
        snippet = t[:length]
        return re.sub(r"\b[A-Za-z0-9]{20,}\b", '<REDACTED>', snippet)
    if isinstance(ai_output, dict):
        logger.info('ai_output is already dict; normalizing and returning')
        return _normalize(ai_output)

    text = '' if ai_output is None else (ai_output if isinstance(ai_output, str) else json.dumps(ai_output, ensure_ascii=False))

    # 尝试直接解析为 JSON
    try:
        parsed = json.loads(text)
        if isinstance(parsed, dict):
            logger.info('direct json.loads succeeded')
            return _normalize(parsed)
        else:
            logger.debug('json.loads returned non-dict: %s', type(parsed).__name__)
            logger.debug('ai_output snippet (masked): %s', _mask_snippet(text))
    except Exception as e:
        logger.debug('direct json.loads failed: %s', e)
        logger.debug('ai_output snippet (masked): %s', _mask_snippet(text))

    # 尝试从文本中提取第一个 JSON 对象子串并解析
    try:
        import re
        m = re.search(r"(\{[\s\S]*\})", text)
        if m:
            cand = m.group(1)
            logger.debug('found JSON substring candidate, attempting parse')
            try:
                parsed = json.loads(cand)
                if isinstance(parsed, dict):
                    logger.info('parsed JSON substring successfully')
                    return _normalize(parsed)
                else:
                    logger.debug('parsed substring is not dict')
                    logger.debug('json substring snippet (masked): %s', _mask_snippet(cand))
            except Exception as e:
                logger.debug('parsing JSON substring failed: %s', e)
                logger.debug('json substring snippet (masked): %s', _mask_snippet(cand))
        else:
            logger.debug('no JSON substring found in text')
    except Exception as e:
        logger.debug('substring extraction failed: %s', e)

    # 最后手段：如果有可用模型，请求模型将输出强制转换为严格的 JSON
    if call_ai:
        try:
            import re as _re
            # strip common code fences and markdown wrappers to give model cleaner input
            cleaned_for_coercion = _re.sub(r'```(?:json)?\s*([\s\S]*?)```', r'\1', text)
            cleaned_for_coercion = cleaned_for_coercion.strip()
        except Exception:
            cleaned_for_coercion = text

        # Primary conversion prompt (clear instructions but allow model to respond)
        example_schema = json.dumps({'core_news': [], '技术前沿': [], '产业动态': [], '政策法规': [], '应用实例': []}, ensure_ascii=False)
        conversion_prompt = (
            "请把下面的文本严格转换为可解析的JSON，返回一个包含五个键的字典：\n"
            f"{example_schema}\n"
            "每个键对应一个数组，数组内每个元素为字符串。不要添加其他文字，请严格只返回 JSON 对象（仅文本，不要代码块或说明）。\n\n"
            f"输入文本：\n{cleaned_for_coercion}\n"
        )

        logger.info('attempting conversion via call_ai (primary) to coerce output to strict JSON')
        try:
            conv = call_ai(model, conversion_prompt)
        except Exception as e:
            logger.error('conversion call_ai (primary) failed: %s', e)
            conv = None

        def _try_parse_conv(candidate):
            try:
                if isinstance(candidate, dict):
                    return candidate
                if not isinstance(candidate, str):
                    return None
                # attempt direct parse
                try:
                    return json.loads(candidate)
                except Exception:
                    # try to extract JSON substring
                    m2 = _re.search(r"(\{[\s\S]*\})", candidate)
                    if m2:
                        try:
                            return json.loads(m2.group(1))
                        except Exception:
                            return None
                    return None
            except Exception:
                return None

        parsed_conv = _try_parse_conv(conv)
        if isinstance(parsed_conv, dict):
            logger.info('conversion (primary) produced dict; normalizing and returning')
            return _normalize(parsed_conv)

        # Secondary stricter coercion: insist only JSON, provide fallback empty schema if necessary
        strict_prompt = (
            "重要：你必须严格且仅返回一个合法的 JSON 对象，并且不要添加任何解释性文字或代码块。"
            " 输出必须是一个 JSON 对象（不允许其他文本）。如果无法从输入中提取内容，请返回空的五键字典。\n\n"
            f"示例输出 (务必只返回类似的 JSON 且不要换行其他内容)：\n{example_schema}\n\n"
            f"输入文本：\n{cleaned_for_coercion}\n\n只返回 JSON："
        )

        logger.info('attempting conversion via call_ai (strict) to coerce output to strict JSON')
        try:
            conv2 = call_ai(model, strict_prompt)
        except Exception as e:
            logger.error('conversion call_ai (strict) failed: %s', e)
            conv2 = None

        parsed_conv2 = _try_parse_conv(conv2)
        if isinstance(parsed_conv2, dict):
            logger.info('conversion (strict) produced dict; normalizing and returning')
            return _normalize(parsed_conv2)

        # if both attempts fail, log masked snippets for debugging and continue to return empty schema
        try:
            logger.warning('conversion attempts failed; primary snippet: %s', _mask_snippet(conv))
            logger.warning('conversion attempts failed; strict snippet: %s', _mask_snippet(conv2))
        except Exception:
            pass

    logger.warning('ensure_structured_ai_response 退化为空的 schema')
    return {k: [] for k in expected_keys}


def sanitize_ai_content(ai_content: Dict[str, Any]) -> Dict[str, Any]:
    """Remove items that look like error payloads or connection/SSL traces
    to avoid writing raw exception dicts into DOCX. Returns a sanitized copy.
    """
    if not isinstance(ai_content, dict):
        return ai_content
    out = {}
    import re

    # common error-like patterns to filter out entire items
    error_patterns = [r"call_deepseek_request_failed", r"HTTPSConnectionPool", r"SSLError", r"Traceback", r"error':", r"Exception"]

    # cleaning regexes for string normalization
    DEEPSEEK_ERROR_RE = re.compile(r"\{\'error\':\s*'call_deepseek_request_failed'.*?\}", re.S)
    CONTROL_CHAR_RE = re.compile(r"[\x00-\x1f\x7f-\x9f]")
    MULTI_EMPTY_RE = re.compile(r"\n{3,}")
    FULLWIDTH_DIGIT_RE = re.compile(r"[０-９]")

    def _fw2hw(m):
        return chr(ord(m.group(0)) - 0xFF10 + ord('0'))

    def _clean_str(s: str) -> str:
        try:
            if s is None:
                return ''
            t = str(s)
            # remove explicit deepseek error dumps
            t = DEEPSEEK_ERROR_RE.sub('', t)
            # remove C0 control chars
            t = CONTROL_CHAR_RE.sub('', t)
            # convert fullwidth digits to ASCII
            t = FULLWIDTH_DIGIT_RE.sub(lambda m: _fw2hw(m), t)
            # collapse excessive blank lines
            t = MULTI_EMPTY_RE.sub('\n\n', t)
            # trim whitespace on each line and remove empty lines
            lines = [ln.strip() for ln in t.splitlines()]
            lines = [ln for ln in lines if ln]
            t = '\n'.join(lines).strip()
            return t
        except Exception:
            return str(s or '')

    def _clean_obj(obj):
        # recursively clean dict/list/str
        if isinstance(obj, str):
            return _clean_str(obj)
        if isinstance(obj, dict):
            new = {}
            for kk, vv in obj.items():
                # skip obvious error keys entirely
                if kk and isinstance(kk, str) and kk.lower() in ('error', 'details', 'traceback', 'exception'):
                    continue
                new_v = _clean_obj(vv)
                # drop empty strings
                if new_v == '' or new_v == [] or new_v == {}:
                    continue
                new[kk] = new_v
            return new
        if isinstance(obj, list):
            res = []
            for e in obj:
                ce = _clean_obj(e)
                if ce == '' or ce == [] or ce == {}:
                    continue
                res.append(ce)
            return res
        # other types -> stringify
        try:
            return _clean_str(str(obj))
        except Exception:
            return str(obj)

    for k, items in ai_content.items():
        new_items = []
        if not isinstance(items, list):
            items = [items]
        for it in items:
            try:
                # if it's a dict that contains obvious error keys, skip
                if isinstance(it, dict):
                    # if dict contains error-like keys/values, skip whole dict
                    keys_lower = {str(x).lower() for x in it.keys()}
                    if keys_lower & {'error', 'details', 'traceback', 'exception'}:
                        continue
                    cleaned = _clean_obj(it)
                    if cleaned:
                        new_items.append(cleaned)
                    continue

                s = _clean_str(it)
                if not s:
                    continue
                # skip lines that match error-like patterns
                skip = False
                for p in error_patterns:
                    if re.search(p, s, flags=re.I):
                        skip = True
                        break
                if skip:
                    continue
                # drop isolated placeholders like '(视频片段)'
                if re.match(r'^[\(\[【].{0,80}[\)\]】]$', s):
                    continue
                new_items.append(s)
            except Exception:
                continue
        out[k] = new_items
    return out


# ----------------- 轻量级数据库助手（当 Qdrant 不可用时使用 SQLite 回退） -----------------
DB_PATH_CANDIDATES = [
    os.path.join(BASE_DIR, 'prompt_templates.db'),
    os.path.join(BASE_DIR, 'instance', 'prompt_templates.db'),
    os.path.join(PARENT, 'instance', 'prompt_templates.db'),
]

def _choose_db_path():
    for p in DB_PATH_CANDIDATES:
        if os.path.exists(p) and os.path.getsize(p) > 0:
            return p
    # 回退：使用模块本地文件（可能为空）
    return os.path.join(BASE_DIR, 'prompt_templates.db')

_DB_PATH = _choose_db_path()

def get_db_conn():
    # 每次调用创建 sqlite3 连接（对本地开发简单且安全）
    conn = sqlite3.connect(_DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def _get_qdrant_client():
    """如果可用且配置了 QDRANT_URL，则返回 QdrantClient，否则返回 None。"""
    if QdrantClient is None:
        logger.debug('qdrant-client not available')
        return None
    url = os.environ.get('QDRANT_URL', 'http://localhost:6333')
    api_key = os.environ.get('QDRANT_API_KEY')
    try:
        client = QdrantClient(url=url, api_key=api_key)
        return client
    except Exception as e:
        logger.warning('failed to create QdrantClient: %s', e)
        return None


def _embed_texts_local(texts: List[str], model_name: str = 'all-MiniLM-L6-v2'):
    """使用本地 sentence-transformers 模型对文本进行嵌入，返回向量列表。"""
    if SentenceTransformer is None:
        raise RuntimeError('sentence-transformers not installed')
    model = SentenceTransformer(model_name)
    embs = model.encode(texts, show_progress_bar=False, convert_to_numpy=True)
    return embs


def rag_fetch_materials(query: str, top_k: int = 3, collection: str = 'collected_articles') -> List[Dict[str, Any]]:
    """在 Qdrant 中搜索查询并返回文章字典列表（包含 title、content、date、score）。

    若缺少依赖或 Qdrant 不可用，会抛出 RuntimeError。
    """
    client = _get_qdrant_client()
    if client is None:
        raise RuntimeError('qdrant client not available')

    # 计算查询向量：仅当显式设置 DEEPSEEK_EMBED_URL 时优先使用 DeepSeek 的外部嵌入服务
    qvec = None
    if os.environ.get('DEEPSEEK_EMBED_URL'):
        try:
            # 动态导入以确保 instance/config.json 已被加载到环境变量
            try:
                from wechat_report_agent.ai_api import call_deepseek_embed as _call_deepseek_embed
            except Exception:
                _call_deepseek_embed = None

            if _call_deepseek_embed is not None:
                logger.info('DEEPSEEK_EMBED_URL present: attempting to use DeepSeek embed API for RAG')
                emb = _call_deepseek_embed(query)
                qvec = emb if isinstance(emb, list) else list(emb)
        except Exception as e:
            logger.warning('DeepSeek embed (DEEPSEEK_EMBED_URL) failed: %s, falling back to local model', e)

    if qvec is None:
        if SentenceTransformer is None:
            raise RuntimeError('sentence-transformers not available for local embedding')
        model_name = os.environ.get('RAG_EMBED_MODEL', 'all-MiniLM-L6-v2')
        model = SentenceTransformer(model_name)
        qvec = model.encode(query, convert_to_numpy=True).tolist()

    try:
        hits = client.search(collection_name=collection, query_vector=qvec, limit=top_k)
    except Exception as e:
        logger.warning('qdrant search failed: %s', e)
        return []

    results = []
    for h in hits:
        # h 类似 ScoredPoint，常含 .payload、.id、.score 等属性
        payload = getattr(h, 'payload', {}) or {}
        title = payload.get('title') or payload.get('name') or ''
        content = payload.get('summary') or payload.get('content') or ''
        date = payload.get('date') or payload.get('create_time') or ''
        score = getattr(h, 'score', None)
        results.append({'id': getattr(h, 'id', None), 'title': title, 'content': content, 'date': date, 'score': score})
    return results

def init_db_if_needed():
    # 确保所需的表存在；此函数可安全重复调用
    conn = get_db_conn()
    c = conn.cursor()
    # templates 表：如旧版使用 'name' 作为列名，我们仍然要兼容
    c.execute('''
    CREATE TABLE IF NOT EXISTS prompt_template (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        title TEXT,
        content TEXT
    );
    ''')
    # 如果现有数据库使用旧列名 'name'，尝试平滑迁移：
    try:
        # 查询表的列信息
        c.execute("PRAGMA table_info(prompt_template)")
        cols = [r[1] for r in c.fetchall()]
        if 'title' not in cols and 'name' in cols:
            logger.info("prompt_template 表缺失 'title' 列，检测到旧列 'name'，开始添加并迁移数据")
            # 添加新列
            c.execute("ALTER TABLE prompt_template ADD COLUMN title TEXT")
            # 将旧的 name 值复制到 title（若 name 存在且 title 为空）
            c.execute("UPDATE prompt_template SET title = name WHERE title IS NULL OR title = ''")
            conn.commit()
            logger.info("prompt_template 表列迁移完成：name -> title")
    except Exception as e:
        logger.exception('尝试迁移 prompt_template 表结构时出错：%s', e)
    # 若缺失则创建 collected_article 和 collect_task 的最小定义
    c.execute('''
    CREATE TABLE IF NOT EXISTS collected_article (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        task_name TEXT,
        title TEXT,
        date TEXT,
        content TEXT
    );
    ''')
    c.execute('''
    CREATE TABLE IF NOT EXISTS collect_task (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT,
        meta TEXT
    );
    ''')
    conn.commit()
    conn.close()

init_db_if_needed()

def row_to_dict(row: sqlite3.Row) -> Dict[str, Any]:
    if row is None:
        return {}
    # 规范化常见列名差异，保证前端获得一致的键名
    d = {k: row[k] for k in row.keys()}
    # 将 'name' 映射为 'title'（用于模板或任务名称）
    if 'name' in d and 'title' not in d:
        d['title'] = d['name']
    # 保留 create_time、summary、source 等字段
    return d

# ----------------- 前端期望的 API 端点 -----------------


@app.route('/api/validate-env', methods=['GET'])
def validate_env():
    required_keys = ['DEEPSEEK_API_KEY', 'DOUBAO_API_KEY', 'ZHIPUAI_API_KEY']
    missing = [k for k in required_keys if not os.environ.get(k)]
    if missing:
        return jsonify({'status': 'error', 'missing_keys': missing}), 400
    return jsonify({'status': 'success'})


@app.route('/api/prompt_templates', methods=['GET', 'POST'])
def api_prompt_templates():
    if request.method == 'GET':
        conn = get_db_conn()
        c = conn.cursor()
        # 返回时兼容旧的 'name' 列与新的 'title' 列，优先使用 title
        c.execute("SELECT id, COALESCE(title, name) as title, content FROM prompt_template ORDER BY id")
        rows = [row_to_dict(r) for r in c.fetchall()]
        conn.close()
        return jsonify(rows)

    # POST -> create
    data = request.json or {}
    title = data.get('title') or data.get('topic') or data.get('name') or 'untitled'
    content = data.get('content') or data.get('template') or ''
    conn = get_db_conn()
    c = conn.cursor()
    # 根据实际表结构决定如何插入：若存在旧列 'name'，同时写入 name 以满足 NOT NULL/UNIQUE 约束
    try:
        c.execute("PRAGMA table_info(prompt_template)")
        cols = [r[1] for r in c.fetchall()]
    except Exception:
        cols = []
    if 'name' in cols:
        c.execute('INSERT INTO prompt_template (name, title, content) VALUES (?, ?, ?)', (title, title, content))
    else:
        c.execute('INSERT INTO prompt_template (title, content) VALUES (?, ?)', (title, content))
    conn.commit()
    new_id = c.lastrowid
    conn.close()
    return jsonify({'id': new_id, 'title': title, 'content': content}), 201


@app.route('/api/prompt_templates/<int:tpl_id>', methods=['PUT', 'DELETE'])
@require_token
def api_prompt_template_modify(tpl_id):
    if request.method == 'DELETE':
        conn = get_db_conn()
        c = conn.cursor()
        c.execute('DELETE FROM prompt_template WHERE id=?', (tpl_id,))
        conn.commit()
        conn.close()
        return jsonify({'status': 'ok', 'deleted_id': tpl_id})

    # PUT -> update
    data = request.json or {}
    title = data.get('title') or data.get('topic')
    content = data.get('content') or data.get('template')
    conn = get_db_conn()
    c = conn.cursor()
    # build dynamic update
    sets = []
    vals = []
    if title is not None:
        sets.append('title=?')
        vals.append(title)
    if content is not None:
        sets.append('content=?')
        vals.append(content)
    if not sets:
        conn.close()
        return jsonify({'status': 'error', 'error': 'no_fields_provided'}), 400
    vals.append(tpl_id)
    # 如果表同时有 name 列，尽量同步更新 name 与 title（保持兼容）
    try:
        c.execute("PRAGMA table_info(prompt_template)")
        cols = [r[1] for r in c.fetchall()]
    except Exception:
        cols = []
    if 'name' in cols and title is not None:
        # 在更新语句中同时设置 name
        sets_with_name = sets.copy()
        # 在 sets_with_name 中加入 name=? 索引在 title 后面
        # 注意：vals 顺序需与 sets 顺序一致
        sets_with_name.insert(0, 'name=?')
        vals = [title] + vals
        c.execute(f"UPDATE prompt_template SET {', '.join(sets_with_name)} WHERE id=?", tuple(vals))
    else:
        c.execute(f"UPDATE prompt_template SET {', '.join(sets)} WHERE id=?", tuple(vals))
    conn.commit()
    # return updated row
    c.execute('SELECT id, COALESCE(title, name) as title, content FROM prompt_template WHERE id=?', (tpl_id,))
    row = c.fetchone()
    conn.close()
    return jsonify(row_to_dict(row))


@app.route('/api/collect/result', methods=['GET'])
def api_collect_result():
    # optional task filter (we don't have a strict FK in collected_article, so do a simple text filter)
    task = request.args.get('task')
    conn = get_db_conn()
    c = conn.cursor()
    # select existing columns; map to front-end friendly keys via row_to_dict
    q = 'SELECT id, title, content, date, summary, source, create_time FROM collected_article'
    params = ()
    if task:
        q += " WHERE title LIKE ? OR content LIKE ?"
        like = f"%{task}%"
        params = (like, like)
    q += ' ORDER BY id DESC'
    c.execute(q, params)
    rows = [row_to_dict(r) for r in c.fetchall()]
    conn.close()
    return jsonify(rows)


@app.route('/api/generate-report', methods=['POST'])
@require_token
def api_generate_report():
    data = request.json or {}
    prompt = data.get('prompt') or data.get('taskName') or '请基于以下材料生成五类要点'
    model = data.get('model') or 'DeepSeek-R1'
    # if client provides 'material' (list of article ids), we can include that
    material = data.get('material') or []

    # If USE_QDRANT enabled, perform a vector retrieval to supplement material
    use_qdrant = os.environ.get('USE_QDRANT', '').lower() in ('1', 'true', 'yes')
    rag_results = []
    if use_qdrant and not material:
        try:
            top_k = int(os.environ.get('RAG_TOPK', '3'))
        except Exception:
            top_k = 3
        try:
            rag_results = rag_fetch_materials(prompt, top_k=top_k)
            if rag_results:
                logger.info('RAG 从 Qdrant 获取到 %d 条材料', len(rag_results))
        except Exception as e:
            logger.warning('RAG fetch failed: %s', e)

    # prepare combined material text if ids provided
    combined_material = ''
    if material:
        try:
            ids = [int(i) for i in material]
            conn = get_db_conn()
            c = conn.cursor()
            q = f"SELECT id, title, content, date, summary, source, create_time FROM collected_article WHERE id IN ({','.join(['?']*len(ids))})"
            c.execute(q, tuple(ids))
            parts = []
            for r in c.fetchall():
                title = r['title'] if 'title' in r.keys() else r.get('name', '')
                date = r['date'] if 'date' in r.keys() else r.get('create_time', '')
                content = r['content'] if 'content' in r.keys() else ''
                parts.append(f"{title}\n{date}\n{content}")
            combined_material = '\n\n'.join(parts)
            conn.close()
        except Exception:
            combined_material = ''

        # if rag_results found and no explicit material by id, include them
        if not combined_material and rag_results:
            parts = []
            for r in rag_results:
                parts.append(f"{r.get('title','')}\n{r.get('date','')}\n{r.get('content','')}")
            combined_material = '\n\n'.join(parts)

    # build prompt including material; if no material provided, add a clear fallback
    fallback_note = (
        "\n\n注意：如果未提供任何材料，请基于公开信息或常识为每个栏目各生成 1-3 条简短要点，"
        "每条为一句话。返回内容必须填充到 JSON 的对应数组中，不要返回除 JSON 之外的任何多余文本。"
    )

    full_prompt = prompt
    if combined_material:
        full_prompt = f"{prompt}\n\n材料：\n{combined_material}"
    else:
        # no material -> instruct model to synthesize illustrative example items
        full_prompt = f"{prompt}{fallback_note}"

    ai_raw = None
    if call_ai:
        try:
            # Build a strict schema and instructions to force AI to use only the provided materials
            schema_example = {
                'core_news': [],
                '技术前沿': [],
                '产业动态': [],
                '政策法规': [],
                '应用实例': []
            }
            schema_text = (
                "请严格返回 JSON，格式示例：\n" + json.dumps(schema_example, ensure_ascii=False, indent=2) + "\n\n"
                "约束：\n"
                "1) 必须仅基于下方提供的材料（'材料' 字段），不得补充外部信息或凭空臆造事实；\n"
                "2) core_news：只返回若干短要点（每条一行，简短句子），不要为 core_news 提供正文或来源；\n"
                "3) 对于每个非 core 的栏目（技术前沿/产业动态/政策法规/应用实例），返回一个对象数组，\n"
                "   每个对象包含 'title'（小标题，1-10 字）和 'summary'（基于材料的正文描述，50-300 字），\n"
                "   summary 必须严格依据材料并在文末标注来源片段（用字段 'source_fragment' 表示）；\n"
                "4) 如果材料不足以支撑某条要点，请返回空数组或在对应位置返回 {\n"
                "   \"title\": \"(材料不足)\", \"summary\": \"(材料不足)\", \"source_fragment\": \"\"\n"
                "   }；\n"
                "5) 请不要返回任何解释、注释或额外文本——仅返回纯 JSON。\n\n"
            )

            payload = schema_text + "材料：\n" + (combined_material or '') + "\n\n请严格按照示例 JSON 输出。"
            ai_raw = call_ai(model, payload)
        except Exception as e:
            logger.exception('call_ai 在 /api/generate-report 中调用失败: %s', e)
            try:
                logger.debug('ai_raw (partial): %s', str(ai_raw)[:200])
            except Exception:
                pass
            ai_raw = None

    ai_content = ensure_structured_ai_response(model, ai_raw)
    # save for debug endpoint
    try:
        _last_ai_debug['ai_raw'] = ai_raw
        _last_ai_debug['ai_content'] = ai_content
        _last_ai_debug['model'] = model
    except Exception:
        pass

    # sanitize ai_content to remove error payloads
    try:
        ai_content = sanitize_ai_content(ai_content)
    except Exception:
        pass

    # prepare normalized rendering content and counts for downstream use
    rendered_for_counts = normalize_ai_content_for_render(ai_content)
    try:
        ai_counts = {k: len(v) for k, v in rendered_for_counts.items()}
    except Exception:
        ai_counts = {}

    # --- enrich summaries: for non-core sections, try to attach a related snippet under each title
    def _synthesize_fragment_by_ai(title, material_text, user_prompt=None, user_prompt_title=None):
        # Strictly use the user's prompt (if provided) to drive per-item generation.
        # If user provided a promptTitle, prefer that; otherwise use prompt. If neither, fall back to internal instr.
        if user_prompt_title:
            base = user_prompt_title
        elif user_prompt:
            base = user_prompt
        else:
            base = (
                "请作为一位在人工智能领域的权威专家，根据给定的标题撰写一段不超过120字的新闻式正文片段。\n"
                "片段应与标题高度相关、精准且可直接作为报告正文引用，仅返回正文内容，不包含额外说明。\n\n"
            )

        payload = f"{base}\n标题：{title}\n材料（可选）：{material_text or ''}\n请直接给出片段："
        try:
            if call_ai:
                res = call_ai(model, payload)
                if isinstance(res, dict):
                    for k in ('text','result','content'):
                        if k in res:
                            return str(res[k])
                    return str(res)
                return str(res)
        except Exception:
            return None
        return None

    try:
        # Only synthesize/overwrite per-item summaries when client explicitly requests it
        do_synthesize = bool(data.get('synthesize_snippets') or data.get('synthesize') or False)
        cm_text = data.get('combined_material') or ''
        if do_synthesize:
            for sec in ("技术前沿", "产业动态", "政策法规", "应用实例"):
                items = rendered_for_counts.get(sec, [])
                for idx, item in enumerate(items):
                    # item is a dict {'title','summary'}
                    # use user's prompt from request data to strictly drive generation
                    user_prompt = data.get('prompt') or data.get('promptTitle') or None
                    user_prompt_title = data.get('promptTitle') if data.get('promptTitle') else None
                    snippet = _synthesize_fragment_by_ai(item.get('title',''), cm_text, user_prompt=user_prompt, user_prompt_title=user_prompt_title)
                    if snippet:
                        rendered_for_counts[sec][idx]['summary'] = snippet
        else:
            # preserve original summaries from AI output (do not auto-overwrite)
            pass
    except Exception:
        logger.exception('synthesize snippets failed')

    # if client requested file download
    if data.get('download'):
        now = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        out = os.path.join(tempfile.gettempdir(), f'report_{now}.docx')
        try:
            rendered = rendered_for_counts
            # ensure each non-core section has at least one placeholder if empty
            for sec in ("技术前沿", "产业动态", "政策法规", "应用实例"):
                if not rendered.get(sec):
                    rendered[sec] = [{'title': '(暂无可用条目)', 'summary': '（该条暂无详细摘要 — 请检查 AI 输出或后端日志）'}]
            # If client provided a template_id (from front-end selection), try to render using that template
            tpl_id = data.get('template_id') or data.get('templateId') or data.get('tpl_id')
            tpl_path = None
            if tpl_id is not None:
                try:
                    conn = get_db_conn()
                    c = conn.cursor()
                    c.execute('SELECT id, COALESCE(title, name) as title, content FROM prompt_template WHERE id=?', (int(tpl_id),))
                    row = c.fetchone()
                    conn.close()
                    if row:
                        # content may be a path to a .docx template or inline template data
                        tpl_path = row['content'] if 'content' in row.keys() else row.get('content')
                except Exception:
                    tpl_path = None

            # If tpl_path appears to be a local .docx file and render_report is available, use it
            used_docx_template = False
            if tpl_path and isinstance(tpl_path, str) and tpl_path.lower().endswith('.docx') and render_report:
                try:
                    # If tpl_path is relative, make it relative to BASE_DIR
                    if not os.path.isabs(tpl_path):
                        tpl_path = os.path.join(BASE_DIR, tpl_path)
                    if os.path.exists(tpl_path):
                        # render_report(template_path, output_path, context)
                        context = {
                            'title': data.get('title') or 'AI 生成报告',
                            'issue': data.get('issue') or '',
                            'org': data.get('org') or '',
                            'date': datetime.date.today().strftime('%Y-%m-%d'),
                            'core_titles': '\n'.join(rendered.get('core_news', [])),
                            'content': rendered
                        }
                        render_report(tpl_path, out, context)
                        used_docx_template = True
                except Exception:
                    used_docx_template = False

            # If tpl_path is present but not a .docx path, treat it as an inline docxtpl template
            # and attempt to render it by creating a temporary .docx base, then using DocxTemplate.
            if not used_docx_template and tpl_path and isinstance(tpl_path, str) and not tpl_path.lower().endswith('.docx'):
                try:
                    # lazy import to avoid hard dependency at module import time
                    from docxtpl import DocxTemplate
                    from docx import Document as _Docx
                    # build a temporary base docx containing the template text
                    base_tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
                    base_tmp.close()
                    doc = _Docx()
                    # preserve blank lines: create paragraphs for each line
                    for line in tpl_path.splitlines() or ['']:
                        p = doc.add_paragraph()
                        p.add_run(line)
                    doc.save(base_tmp.name)

                    # render the template using DocxTemplate
                    dt = DocxTemplate(base_tmp.name)
                    context = {
                        'title': data.get('title') or 'AI 生成报告',
                        'issue': data.get('issue') or '',
                        'org': data.get('org') or '',
                        'date': datetime.date.today().strftime('%Y-%m-%d'),
                        'core_titles': '\n'.join(rendered.get('core_news', [])),
                        'content': rendered
                    }
                    dt.render(context)
                    dt.save(out)
                    used_docx_template = True
                except Exception:
                    # if anything fails (missing packages or render error), fall back later
                    used_docx_template = False
                finally:
                    try:
                        if 'base_tmp' in locals() and os.path.exists(base_tmp.name):
                            os.remove(base_tmp.name)
                    except Exception:
                        pass

            if not used_docx_template:
                if generate_ai_report:
                    generate_ai_report(rendered, out, combined_material=combined_material)
                else:
                    simple_generate_docx(rendered, out)
            resp = send_file(out, as_attachment=True, download_name=os.path.basename(out))
            # attach AI counts as a header so frontend can read per-section counts without parsing DOCX
            try:
                import json as _json
                # Header must be ASCII-safe; force ASCII encoding for header value
                resp.headers['X-AI-COUNTS'] = _json.dumps(ai_counts, ensure_ascii=True)
            except Exception:
                pass
            return resp
        except Exception as e:
            # 记录完整错误与可能的 ai_raw 片段
            logger.exception('生成报告失败')
            try:
                logger.debug('ai_content (partial): %s', str(ai_content)[:800])
            except Exception:
                pass
            return jsonify({'status': 'error', 'error': str(e)}), 500

    return jsonify({'status': 'ok', 'ai_content': ai_content, 'ai_counts': ai_counts})


@app.route('/api/clear-all-data', methods=['POST'])
@require_token
def api_clear_all_data():
    # wipe collected_article and collect_task, and optionally prompt_template
    conn = get_db_conn()
    c = conn.cursor()
    c.execute('DELETE FROM collected_article')
    c.execute('DELETE FROM collect_task')
    # do not delete prompt_template by default unless requested
    if request.json and request.json.get('wipe_templates'):
        c.execute('DELETE FROM prompt_template')
    conn.commit()
    conn.close()
    return jsonify({'status': 'ok'})


@app.route('/api/test-models', methods=['POST'])
@require_token
def test_models():
    data = request.json or {}
    prompt = data.get('prompt', '测试提示词')
    models = ['DeepSeek-R1', '豆包', '智谱AI']
    results = {}
    for m in models:
        try:
            if call_ai:
                results[m] = call_ai(m, prompt)
            else:
                results[m] = 'mock-response'
        except Exception as e:
            results[m] = f'error: {e}'
    return jsonify(results)


@app.route('/api/generate-report-debug', methods=['GET', 'POST'])
@require_token
def generate_report_debug():
    data = request.json or {}
    prompt = data.get('prompt', '示例：请基于以下材料生成五类要点')
    model = 'DeepSeek-R1'

    schema = {
        "core_news": [],
        "技术前沿": [],
        "产业动态": [],
        "政策法规": [],
        "应用实例": []
    }

    # 调用 AI 并确保输出为结构化五键 JSON
    ai_raw = None
    if call_ai:
        try:
            ai_raw = call_ai(model, f"请严格返回JSON:\n{json.dumps(schema, ensure_ascii=False)}\n\n{prompt}")
        except Exception as e:
            logger.warning('call_ai 调用失败: %s', e)
            ai_raw = None

    ai_content = ensure_structured_ai_response(model, ai_raw)
    try:
        _last_ai_debug['ai_raw'] = ai_raw
        _last_ai_debug['ai_content'] = ai_content
        _last_ai_debug['model'] = model
    except Exception:
        pass

    now = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
    out = os.path.join(tempfile.gettempdir(), f'debug_report_{now}.docx')

    rendered = normalize_ai_content_for_render(ai_content)
    try:
        if generate_ai_report:
            generate_ai_report(rendered, out, combined_material=None)
        else:
            simple_generate_docx(rendered, out)
        resp = send_file(out, as_attachment=True, download_name=os.path.basename(out))
        # attach ai_counts header to debug send_file response as well
        try:
            import json as _json
            # Ensure header contains only ASCII characters to avoid send_header UnicodeEncodeError
            resp.headers['X-AI-COUNTS'] = _json.dumps({k: len(v) for k, v in rendered.items()}, ensure_ascii=True)
        except Exception:
            pass
        return resp
    except Exception as e:
        logger.error('debug generate send_file failed: %s', e)
        return jsonify({'status': 'error', 'error': str(e)}), 500


@app.route('/api/debug-ai-output', methods=['GET'])
@require_token
def api_debug_ai_output():
    """Return the last AI raw output and normalized content for debugging.

    Only available when authenticated with REPORT_API_TOKEN (or DEV_AUTH_DISABLED local bypass).
    """
    return jsonify({'status': 'ok', 'last': _last_ai_debug})


@app.route('/system/menu/my/1.6.1', methods=['GET'])
@require_token
def api_system_menu_my():
    """Return a menu structure compatible with the SCUI frontend.

    The frontend expects menu items with fields like:
      - name, path, component, meta: { title, icon, ... }, children: []
    """
    # Build a menu array that matches the SCUI dynamic router expectations
    menu_array = [
        {
            'name': 'dashboard',
            'path': '/dashboard',
            'component': 'dashboard/index',
            'meta': {'icon': 'ElIconDataAnalysis', 'title': '首页'},
            'children': [
                {'name': 'collect_task_content', 'path': '/dashboard/collect-task-content', 'component': 'dashboard/collect-task-content', 'meta': {'title': '采集内容', 'icon': 'ElIconNotebook'}},
                                {'name': 'prompt_word_setting', 'path': '/dashboard/prompt-word-setting', 'component': 'dashboard/prompt-word-setting', 'meta': {'title': '提示词管理', 'icon': 'ElIconManagement'}},
                # output_template_settings removed per request: UI/page deleted
                {'name': 'report_generation', 'path': '/dashboard/report-generation', 'component': 'dashboard/report-generation', 'meta': {'title': '报告生成', 'icon': 'ElIconEdit'}},
                {'name': 'generated_report', 'path': '/dashboard/generated-report', 'component': 'dashboard/generated-report', 'meta': {'title': '已生成报告', 'icon': 'ElIconDocument'}},
                {'name': 'automatic_release_settings', 'path': '/dashboard/automatic-release-settings', 'component': 'dashboard/automatic-release-settings', 'meta': {'title': '自动发送设置', 'icon': 'ElIconCheck'}}
            ]
        }
    ]

    # sanitize strings in menu to remove stray control/newline characters
    def _sanitize_value(v):
        if not isinstance(v, str):
            return v
        # remove control characters and normalize whitespace
        cleaned = ''.join(ch for ch in v if ch >= ' ')
        s = ' '.join(cleaned.split())
        return s

    def sanitize_menu_strings(node):
        if isinstance(node, dict):
            for k, val in list(node.items()):
                if isinstance(val, (dict, list)):
                    sanitize_menu_strings(val)
                else:
                    # sanitize component and path specially
                    if k == 'component' and isinstance(val, str):
                        node[k] = _sanitize_value(val).replace('\n', '').replace('\r', '').strip()
                    elif k == 'path' and isinstance(val, str):
                        p = _sanitize_value(val).strip()
                        # strip leading /scui prefix if present
                        if p.startswith('/scui'):
                            p = p.replace('/scui', '', 1) or '/'
                        node[k] = p
                    else:
                        node[k] = _sanitize_value(val)
        elif isinstance(node, list):
            for item in node:
                sanitize_menu_strings(item)

    sanitize_menu_strings(menu_array)

    # final deep-clean: ensure no stray \n, \r, \t and remove any repeated whitespace
    import re
    def deep_clean(node):
        if isinstance(node, dict):
            for k, v in list(node.items()):
                if isinstance(v, (dict, list)):
                    deep_clean(v)
                elif isinstance(v, str):
                    s = re.sub(r'[\r\n\t]+', ' ', v)
                    s = re.sub(r'\s{2,}', ' ', s).strip()
                    if k == 'path' and s.startswith('/scui'):
                        s = s.replace('/scui', '', 1) or '/'
                    node[k] = s
        elif isinstance(node, list):
            for item in node:
                deep_clean(item)

    deep_clean(menu_array)

    # Normalize icon strings: convert legacy class names (e.g. 'el-icon-data-analysis')
    # to registered component names used by the frontend (e.g. 'ElIconDataAnalysis').
    def _pascalize(name: str) -> str:
        parts = [p for p in name.split('-') if p]
        return ''.join(p.capitalize() for p in parts)

    def normalize_icon_string(icon_val):
        if not isinstance(icon_val, str):
            return icon_val
        s = icon_val.strip()
        # handle Element class-style icons like 'el-icon-data-analysis' -> 'ElIconDataAnalysis'
        if s.startswith('el-icon-'):
            core = s[len('el-icon-'):]
            return 'ElIcon' + _pascalize(core)
        # handle project sc icons 'sc-icon-xxx' -> 'ScIconXxx'
        if s.startswith('sc-icon-'):
            core = s[len('sc-icon-'):]
            return 'ScIcon' + _pascalize(core)
        # if already looks like a component name, leave as-is
        return s

    def normalize_menu_icons(node):
        if isinstance(node, dict):
            # update meta.icon if present
            meta = node.get('meta')
            if isinstance(meta, dict) and 'icon' in meta:
                meta['icon'] = normalize_icon_string(meta.get('icon'))
            for v in node.values():
                if isinstance(v, (dict, list)):
                    normalize_menu_icons(v)
        elif isinstance(node, list):
            for item in node:
                normalize_menu_icons(item)

    normalize_menu_icons(menu_array)

    envelope = {
        'code': 200,
        'data': {
            'menu': menu_array,
            'dashboardGrid': [],
            'permissions': ['admin:all']
        },
        'message': 'success'
    }
    return jsonify(envelope)


# Backwards-compatible alias used by frontend code (some clients call /api/system/...)
@app.route('/api/system/menu/my/1.6.1', methods=['GET'])
@require_token
def api_system_menu_my_alias():
    return api_system_menu_my()


# store last client-reported menu for debugging
_last_client_menu = None


@app.route('/api/debug/client-menu', methods=['POST'])
def api_debug_client_menu_post():
    global _last_client_menu
    try:
        payload = request.get_json(force=True)
    except Exception:
        payload = None
    # store both the raw payload and a JSON-string snapshot to avoid serialization issues
    try:
        payload_json = json.dumps(payload, ensure_ascii=False, default=str)
    except Exception:
        payload_json = None
    _last_client_menu = {
        'time': time.time(),
        'payload': payload,
        'payload_json': payload_json
    }
    return jsonify({'status': 'ok'})


@app.route('/api/debug/client-menu', methods=['GET'])
def api_debug_client_menu_get():
    """Return last client debug snapshot (if any)."""
    try:
        if _last_client_menu is None:
            return jsonify({'status': 'empty', 'last': None})
        return jsonify({'status': 'ok', 'last': _last_client_menu})
    except Exception:
        return jsonify({'status': 'error', 'error': 'failed_to_return_snapshot'}), 500


# Dev token endpoints (compatibility for frontend during local development)
@app.route('/token', methods=['POST'])
@app.route('/api/token', methods=['POST'])
def api_token():
    # prefer environment REPORT_API_TOKEN, then instance/config.json, else default '1'
    token = os.environ.get('REPORT_API_TOKEN')
    if not token:
        try:
            cfg_path = os.path.join(os.path.dirname(__file__), 'instance', 'config.json')
            if os.path.exists(cfg_path):
                with open(cfg_path, 'r', encoding='utf-8') as f:
                    cfg = json.load(f)
                    token = cfg.get('REPORT_API_TOKEN')
        except Exception:
            token = None

    if not token:
        token = '1'

    userInfo = {
        'username': 'admin',
        'role': ['admin'],
        'name': 'Developer'
    }

    return jsonify({'code': 200, 'data': {'token': token, 'userInfo': userInfo}})


if __name__ == '__main__':
    app.run(host='127.0.0.1', port=5001, debug=False, use_reloader=False)
