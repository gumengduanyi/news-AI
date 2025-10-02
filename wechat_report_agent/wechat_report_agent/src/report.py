from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cn_font(run, name="宋体", size=12, bold=False, color=None):
    """统一设置中文字体为宋体"""
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_red_line_cell(cell, size=6):
    """给单元格下方加红色横线"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')      # 单线
    bottom.set(qn('w:sz'), str(size))      # 粗细，单位 1/8 磅
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), 'FF0000')    # 红色
    tcBorders.append(bottom)

def build_report(meta, items, output_path):
    doc = Document()

    # ====== 封面部分 ======
    # 大标题
    p1 = doc.add_paragraph()
    run1 = p1.add_run(meta["title"])
    set_cn_font(run1, size=20, bold=True, color=RGBColor(192, 0, 0))
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 期号
    p2 = doc.add_paragraph()
    run2 = p2.add_run(meta["issue"])
    set_cn_font(run2, size=14, bold=True)
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")

    # 左右分布：组织在左，日期在右
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    row = table.rows[0]
    cell1, cell2 = row.cells

    p_left = cell1.paragraphs[0]
    run_left = p_left.add_run(meta["org"])
    set_cn_font(run_left, size=12, bold=True)
    p_left.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p_right = cell2.paragraphs[0]
    run_right = p_right.add_run(meta["date"])
    set_cn_font(run_right, size=12, bold=True)
    p_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # 给整行两个单元格都加红线
    for cell in row.cells:
        add_red_line_cell(cell, size=6)

    doc.add_paragraph("")

    # ====== 本期核心要闻 ======
    p4 = doc.add_paragraph()
    run4 = p4.add_run("本期核心要闻")
    set_cn_font(run4, size=14, bold=True)
    p4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    sections = ["技术前沿", "产业动态", "政策法规", "应用实例"]
    core_titles = []
    for sec in sections:
        sec_items = [
            x for x in items
            if x and x.get("category") == sec
            and x.get("title")
            and not x["title"].startswith("无标题")
            and x.get("summary")
        ]
        for it in sec_items:
            core_titles.append(it["title"])

    for t in core_titles:
        para = doc.add_paragraph()
        run = para.add_run(f"● {t}")
        set_cn_font(run, size=12, bold=False)

    # ====== 正文分章节 ======
    sections = ["技术前沿", "产业动态", "政策法规", "应用实例", "工作建议"]
    for sec in sections:
        sec_items = [
            x for x in items 
            if x and x.get("category") == sec 
            and x.get("title") 
            and not x["title"].startswith("无标题") 
            and x.get("summary")
        ]
        if not sec_items:
            continue  

        # 一级标题
        h = doc.add_paragraph()
        run_h = h.add_run(sec)
        set_cn_font(run_h, size=14, bold=True)

        for idx, it in enumerate(sec_items, 1):
            # 小标题
            p_title = doc.add_paragraph()
            run_t = p_title.add_run(f"{idx}. {it['title']}")
            set_cn_font(run_t, size=12, bold=True)

            # 摘要正文
            p_sum = doc.add_paragraph()
            run_s = p_sum.add_run(it['summary'])
            set_cn_font(run_s, size=11, bold=False)

    doc.save(output_path)
    print(f"[Report] 已生成报告: {output_path}")
