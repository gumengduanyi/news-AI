# -*- coding: utf-8 -*-
"""
示例：用 docxtpl 渲染 Word 模板（支持 {{title}}、{{issue}}、{{org}}、{{date}}、{{core_titles}}、{{content}} 占位符）
"""
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import datetime
import logging
import os
import re
import difflib
import string

logger = logging.getLogger('render_word_report')
if not logger.handlers:
    h = logging.StreamHandler()
    h.setFormatter(logging.Formatter('%(asctime)s %(levelname)s %(message)s'))
    logger.addHandler(h)
logger.setLevel(logging.INFO)

def render_report(template_path, output_path, context):
    tpl = DocxTemplate(template_path)
    tpl.render(context)
    tpl.save(output_path)

def set_cn_font(run, name="FZFSK", size=16, bold=False, color=None):
    """统一设置中文字体为FZFSK，数字和英文为Times New Roman"""
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def set_title_font(run, name="FZFSK", size=28, bold=True, color=None):
    """设置大标题字体为FZFSK，28磅字"""
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def set_subtitle_font(run, name="方正黑体_GBK", size=16, bold=True, color=None):
    """设置小标题字体为方正黑体_GBK，三号字"""
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_red_line_cell(cell, size=6):
    """给单元格下方加红色横线"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')      # 单线
    bottom.set(qn('w:sz'), str(size))      # 粗细，单位 1/8 磅
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), 'FF0000')    # 红色
    tcBorders.append(bottom)

def _truncate_text(text, max_chars=800):
    if not text:
        return text
    s = str(text).strip()
    if len(s) <= max_chars:
        return s
    # try to cut at sentence boundary
    cut = s[:max_chars]
    last = max(cut.rfind('。'), cut.rfind('.'), cut.rfind('!'), cut.rfind('！'), cut.rfind('?'), cut.rfind('？'))
    if last > int(max_chars * 0.6):
        return cut[:last+1] + '…'
    return cut.rstrip() + '…'


def _clean_summary(text, combined_material=None):
    if not text:
        return ''
    s = str(text)
    # remove typical AI markers and headings
    s = re.sub(r'(AI\s*生成：|AI Generated:|相关片段：|来源片段：)\s*', '', s, flags=re.I)
    # remove explicit manual/data-source markers like 【人工】 or similar
    s = re.sub(r'【[^】]{0,60}人工[^】]{0,60}】', '', s)
    s = re.sub(r'\[人工\]', '', s)
    s = re.sub(r'(^|\n)【人工】\s*', '\1', s)
    # remove excessive whitespace
    s = re.sub(r'[ \t\u00A0]{2,}', ' ', s)
    s = s.strip()
    # remove long parenthetical or bracketed literary fragments which often are narrative descriptions
    try:
        # full-width parentheses/brackets like （...） or 【...】
        s = re.sub(r'（[^）]{30,}）', '', s)
        s = re.sub(r'\([^\)]{30,}\)', '', s)
        s = re.sub(r'【[^】]{30,}】', '', s)
        s = re.sub(r'\[[^\]]{30,}\]', '', s)
    except Exception:
        pass
    # Previously we removed sentences that matched combined_material verbatim to avoid paste.
    # That led to summaries being emptied when the AI output was legitimately drawn from the material.
    # Change: do NOT remove verbatim sentences here. Upstream rendering prefers material evidence explicitly
    # via _extract_evidence and will attach a source fragment instead. Keeping verbatim sentences prevents
    # accidental loss of all summary text.
    # (No-op placeholder left for backward-compatibility.)
    # if combined_material and isinstance(combined_material, str):
    #     pass
    # additional heuristic: if the remaining text looks like fiction/narrative (many personal pronouns and commas), drop it
    try:
        # count simple narrative indicators: Chinese pronouns 他/她/我/我们 and punctuation density
        pronouns = len(re.findall(r'[他她我我们你他们她们]', s))
        commas = len(re.findall(r'[，,]', s))
        if pronouns >= 3 and commas >= 4 and len(s) > 80:
            logger.info('Filtered narrative-like summary (pronouns=%d, commas=%d, len=%d)', pronouns, commas, len(s))
            return ''
    except Exception:
        pass
    return s


def _split_sentences(text: str) -> list:
    if not text:
        return []
    t = text.replace('。', '。').replace('！', '。').replace('?', '。').replace('？', '。')
    # normalize newlines to sentence breaks
    t = t.replace('\n', '。')
    sents = [s.strip() for s in t.split('。') if s.strip()]
    return sents


def _extract_evidence(title: str, combined_material: str, max_sentences: int = 2) -> str:
    """从 combined_material 中抽取与 title 最相关的句子作为证据（简单启发式）。

    策略：
    - 优先找包含 title 的句子；
    - 否则找包含 title 中较长子串（>=2字符）的句子；
    - 否则使用序列匹配选出相似度最高的句子（ratio >= 0.45）。
    返回拼接的句子（以 '。' 分隔）。
    """
    if not (title and combined_material):
        return ''
    sents = _split_sentences(combined_material)
    if not sents:
        return ''

    selected = []
    # exact containment
    for s in sents:
        if title in s:
            selected.append(s)
            if len(selected) >= max_sentences:
                return '。'.join(selected)

    # substring matching (longer substrings first)
    # remove common ASCII and Chinese punctuation to build substrings
    punct = string.punctuation + '，。！？；：、【】（）《》“”‘’—…·'  # extend as needed
    try:
        clean_title = re.sub(rf"[{re.escape(punct)}\s]+", '', title)
    except Exception:
        clean_title = re.sub(r'\s+', '', title)
    # for Chinese, check substring lengths
    L = len(clean_title)
    for n in range(min(6, max(2, L)), 1, -1):
        if len(clean_title) < n:
            continue
        found = False
        for i in range(0, len(clean_title) - n + 1):
            sub = clean_title[i:i+n]
            if len(sub) < 2:
                continue
            for s in sents:
                if sub in s and s not in selected:
                    selected.append(s)
                    found = True
                    if len(selected) >= max_sentences:
                        return '。'.join(selected)
        if found:
            break

    # fallback: sequence matching
    ratios = [(difflib.SequenceMatcher(None, title, s).ratio(), s) for s in sents]
    ratios.sort(reverse=True, key=lambda x: x[0])
    for r, s in ratios:
        if r >= 0.45 and s not in selected:
            selected.append(s)
            if len(selected) >= max_sentences:
                break

    return '。'.join(selected)


def normalize_ai_content_for_render(ai_content):
    # ensure dict with expected keys and lists
    expected = ["core_news", "技术前沿", "产业动态", "政策法规", "应用实例"]
    out = {}
    if not isinstance(ai_content, dict):
        return {k: [] for k in expected}
    for k in expected:
        v = ai_content.get(k, [])
        if v is None:
            v = []
        if isinstance(v, list):
            out[k] = v
        else:
            out[k] = [v]
    return out


def set_small_font(run, name="宋体", size=9, italic=False, color=None):
    try:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
        run.font.size = Pt(size)
        run.font.italic = italic
        if color:
            run.font.color.rgb = color
    except Exception:
        pass


def generate_ai_report(ai_content, output_path, combined_material=None, meta=None, include_sources=False):
    doc = Document()
    # log a short summary of ai_content to help debugging when doc has no body
    try:
        counts = {k: len(v) if v is not None else 0 for k, v in ai_content.items()} if isinstance(ai_content, dict) else {}
    except Exception:
        counts = {}
    logger.info('generate_ai_report called; ai_content counts=%s', counts)

    # ====== 封面部分 ======
    # Use meta for cover if provided, otherwise defaults
    meta = meta or {}
    title_text = meta.get('title', '人工智能技术跟踪周报告')
    issue_text = meta.get('issue', '第12期')
    org_text = meta.get('org', '新闻智能体')
    date_text = meta.get('date', datetime.date.today().strftime('%Y-%m-%d'))

    # 大标题
    p1 = doc.add_paragraph()
    run1 = p1.add_run(title_text)
    set_title_font(run1)
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 期号
    p2 = doc.add_paragraph()
    run2 = p2.add_run(issue_text)
    set_cn_font(run2, size=14, bold=True)
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")

    # 左右分布：组织在左，日期在右
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    row = table.rows[0]
    cell1, cell2 = row.cells

    p_left = cell1.paragraphs[0]
    run_left = p_left.add_run(org_text)
    set_cn_font(run_left, size=12, bold=True)
    p_left.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p_right = cell2.paragraphs[0]
    run_right = p_right.add_run(date_text)
    set_cn_font(run_right, size=12, bold=True)
    p_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # 给整行两个单元格都加红线
    for cell in row.cells:
        add_red_line_cell(cell, size=6)

    doc.add_paragraph("")

    # ====== 本期核心要闻 ======
    p4 = doc.add_paragraph()
    run4 = p4.add_run("本期核心要闻")
    set_cn_font(run4, size=14, bold=True)
    p4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # normalize core_news and sections
    norm = normalize_ai_content_for_render(ai_content)

    for item in norm.get("core_news", []):
        para = doc.add_paragraph()
        # use double marker like '● ⚫' to match desired style
        run = para.add_run(f"● ⚫ {item}")
        set_cn_font(run, size=12, bold=False)

    # If no core_news and no other sections, add a visible placeholder so users know AI content was empty
    has_core = bool(ai_content.get('core_news')) if isinstance(ai_content, dict) else False
    has_other = any(bool(ai_content.get(sec)) for sec in (["技术前沿", "产业动态", "政策法规", "应用实例"])) if isinstance(ai_content, dict) else False
    if not has_core and not has_other:
        p_note = doc.add_paragraph()
        run_note = p_note.add_run('（未从 AI 获取到正文内容 — 请检查后端日志）')
        set_cn_font(run_note, size=12, bold=False)

    # ====== 正文分章节 ======
    sections = ["技术前沿", "产业动态", "政策法规", "应用实例"]
    for sec in sections:
        sec_items = norm.get(sec, [])
        if not sec_items:
            continue

        # 一级标题
        h = doc.add_paragraph()
        run_h = h.add_run(sec)
        set_cn_font(run_h, size=14, bold=True)

        for idx, it in enumerate(sec_items, 1):
            if isinstance(it, str):
                it = {'title': it, 'summary': ''}

            # 小标题
            p_title = doc.add_paragraph()
            run_t = p_title.add_run(f"{idx}. {it['title']}")
            set_subtitle_font(run_t)

            # 摘要正文
            # 如果 summary 为空但 title 本身像是完整句子或很长，则把 title 当作摘要写入，避免只显示小标题无正文
            summary = it.get('summary', '') or ''
            title_text = it.get('title', '') or ''
            # heuristics: treat title as summary when it's long or contains punctuation/newline
            if not summary and (len(title_text) > 60 or any(p in title_text for p in ('.', '。', '\n', '，', ';'))):
                summary = title_text
            # 如果仍然没有摘要，插入占位提示以提醒后台可能未生成正文
            if not summary:
                summary = '（该条暂无详细摘要 — 请检查 AI 输出或后端日志）'

            # If combined_material provided, try to extract evidence from it first
            extracted = ''
            if combined_material:
                try:
                    extracted = _extract_evidence(title_text, combined_material, max_sentences=2)
                except Exception:
                    extracted = ''

            if extracted:
                # prefer evidence from material; mark as source
                cleaned = _clean_summary(extracted, combined_material)
                source_fragment = extracted
            else:
                # clean and truncate summary (fallback to AI-produced summary)
                cleaned = _clean_summary(summary, combined_material)
                source_fragment = it.get('source') or it.get('origin') or None
            # if no cleaned content but title is long/contains punctuation, use title as summary
            if not cleaned and (len(title_text) > 60 or any(p in title_text for p in ('.', '。', '\n', '，', ';'))):
                cleaned = title_text
            if not cleaned:
                cleaned = '（该条暂无详细摘要 — 请检查 AI 输出或后端日志）'

            # enforce max chars per item (800 chars)
            cleaned = _truncate_text(cleaned, max_chars=800)

            p_sum = doc.add_paragraph()
            run_s = p_sum.add_run(cleaned)
            set_cn_font(run_s, size=11, bold=False)

            # optionally include source fragment or combined_material snippet as small italic footnote-like line
            if include_sources:
                src = source_fragment
                if not src and combined_material:
                    # include a short prefix of combined_material if available (first 120 chars)
                    src = combined_material.strip()[:120]
                if src:
                    p_src = doc.add_paragraph()
                    run_src = p_src.add_run(f"来源片段：{_truncate_text(src, max_chars=200)}")
                    set_small_font(run_src, size=9, italic=True)
            # 不再插入额外的 AI 生成段落或原文片段，保留 summary 作为正文

    doc.save(output_path)
    print(f"[Report] 已生成报告: {output_path}")

if __name__ == '__main__':
    context = {
        'title': 'AI 新闻周报',
        'issue': '第12期',
        'org': '新闻智能体',
        'date': datetime.date.today().strftime('%Y-%m-%d'),
        'core_titles': '1. 要闻A\n2. 要闻B',
        'content': '这里是正文内容...'
    }
    render_report('report_template.docx', 'output.docx', context)
    print('报告已生成：output.docx')
