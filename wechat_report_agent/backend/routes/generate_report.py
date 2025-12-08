"""Generate-report handlers migrated from prompt_qdrant_api.py.

This module provides incremental replacement for the monolithic route.
It tries to reuse helpers from the original module when available and
falls back to safe behaviors when not.
"""
import os
import json
import logging
import datetime
import tempfile
from flask import jsonify, send_file, Blueprint, request
from wechat_report_agent.backend.db import get_db_conn

logger = logging.getLogger('backend.routes.generate')

from wechat_report_agent.backend.auth import require_token

# Blueprint for report generation endpoints
generate_report_bp = Blueprint('generate_report', __name__)
generate_report_bp = generate_report_bp


def _fetch_materials_by_ids(material_ids):
    if not material_ids:
        return []
    try:
        ids = [int(i) for i in material_ids]
    except Exception:
        ids = []
    if not ids:
        return []
    conn = get_db_conn()
    c = conn.cursor()
    q = f"SELECT id, title, content, date, summary, source, create_time FROM collected_article WHERE id IN ({','.join(['?']*len(ids))})"
    c.execute(q, tuple(ids))
    rows = []
    for r in c.fetchall():
        rows.append({'id': r['id'], 'title': r.get('title') or r.get('name', ''), 'content': r.get('content','')})
    conn.close()
    return rows


# Prefer helpers from the backend package; fall back to conservative defaults.
try:
    from wechat_report_agent.backend.ai_utils import (
        call_ai,
        ensure_structured_ai_response,
        normalize_ai_content_for_render,
        sanitize_ai_content,
        rag_fetch_materials,
    )
except Exception:
    # ensure names exist for import-checks even when module import fails
    call_ai = None
    ensure_structured_ai_response = lambda model, out: {
        'core_news': [], '技术前沿': [], '产业动态': [], '政策法规': [], '应用实例': []
    }
    normalize_ai_content_for_render = lambda x: x or {}
    sanitize_ai_content = lambda x: x
    rag_fetch_materials = None

# docx helpers
try:
    from wechat_report_agent.backend.docx_utils import simple_generate_docx
except Exception:
    simple_generate_docx = None

try:
    from wechat_report_agent.src.render_word_report import generate_ai_report, render_report
except Exception:
    generate_ai_report = None
    render_report = None

try:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
except Exception:
    BASE_DIR = '.'


def _build_combined_material(material_ids, prompt):
    # Try explicit ids first, then RAG if available
    combined = ''
    if material_ids:
        parts = _fetch_materials_by_ids(material_ids)
        combined = '\n\n'.join((p.get('title','') + '\n' + p.get('content','')) for p in parts)
    if not combined and rag_fetch_materials:
        try:
            rag = rag_fetch_materials(prompt, top_k=int(os.environ.get('RAG_TOPK', '3')))
            if rag:
                combined = '\n\n'.join((r.get('title','') + '\n' + r.get('content','')) for r in rag)
        except Exception as e:
            logger.debug('rag fetch failed: %s', e)
    return combined


def _dump_docx_and_respond(rendered, combined_material=None):
    now = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
    out = os.path.join(tempfile.gettempdir(), f'report_{now}.docx')
    # Prefer generate_ai_report if available
    try:
        if generate_ai_report:
            generate_ai_report(rendered, out, combined_material=combined_material)
        elif simple_generate_docx:
            simple_generate_docx(rendered, out)
        else:
            # fall back to writing a tiny placeholder docx using python-docx if available
            try:
                from docx import Document
                doc = Document()
                doc.add_heading('AI 生成报告（占位）', level=1)
                for k, items in (rendered or {}).items():
                    doc.add_heading(str(k), level=2)
                    if not items:
                        doc.add_paragraph('(无内容)')
                    else:
                        for it in items:
                            doc.add_paragraph(str(it))
                doc.save(out)
            except Exception:
                # As last resort, create a text file with .docx extension so frontend can still download
                with open(out, 'w', encoding='utf-8') as f:
                    f.write(json.dumps(rendered or {}, ensure_ascii=False, indent=2))
    except Exception as e:
        logger.exception('failed to render docx: %s', e)
        raise

    resp = send_file(out, as_attachment=True, download_name=os.path.basename(out))
    try:
        import json as _json
        resp.headers['X-AI-COUNTS'] = _json.dumps({k: len(v) for k, v in (rendered or {}).items()}, ensure_ascii=True)
    except Exception:
        pass
    return resp


def _synthesize_fragment_by_ai(title, material_text, model='DeepSeek-R1', user_prompt=None, user_prompt_title=None):
    if user_prompt_title:
        base = user_prompt_title
    elif user_prompt:
        base = user_prompt
    else:
        base = (
            "请作为一位在人工智能领域的权威专家，根据给定的标题撰写一段不超过120字的新闻式正文片段。\n"
            "片段应与标题高度相关、精准且可直接作为报告正文引用，仅返回正文内容，不包含额外说明。\n\n"
        )
    payload = f"{base}\n标题：{title}\n材料（可选）：{material_text or ''}\n请直接给出片段："
    try:
        if call_ai:
            res = call_ai(model, payload)
            if isinstance(res, dict):
                for k in ('text', 'result', 'content'):
                    if k in res:
                        return str(res[k])
                return str(res)
            return str(res)
    except Exception:
        return None
    return None


def handle_generate_report(req):
    data = req.json or {}
    prompt = data.get('prompt') or data.get('taskName') or '请基于以下材料生成五类要点'
    model = data.get('model') or 'DeepSeek-R1'
    material = data.get('material') or []

    combined_material = _build_combined_material(material, prompt)

    ai_raw = None
    ai_content = None

    if call_ai:
        try:
            # Minimal instruction: ask AI to return expected schema
            schema = {
                'core_news': [], '技术前沿': [], '产业动态': [], '政策法规': [], '应用实例': []
            }
            instr = '请严格返回 JSON，键为 core_news, 技术前沿, 产业动态, 政策法规, 应用实例。材料字段如下：\n' + (combined_material or '')
            ai_raw = call_ai(model, instr)
            ai_content = ensure_structured_ai_response(model, ai_raw)
        except Exception as e:
            logger.exception('call_ai failed in migrated handler: %s', e)
            ai_raw = None
            ai_content = ensure_structured_ai_response(model, None)
    else:
        # No AI available -> return empty schema with diagnostics
        ai_content = ensure_structured_ai_response(model, None)

    try:
        ai_content = sanitize_ai_content(ai_content)
    except Exception:
        pass

    rendered = normalize_ai_content_for_render(ai_content)
    try:
        ai_counts = {k: len(v) for k, v in rendered.items()}
    except Exception:
        ai_counts = {}

    # synthesize per-item snippets if requested
    try:
        do_synthesize = bool(data.get('synthesize_snippets') or data.get('synthesize') or False)
        cm_text = data.get('combined_material') or combined_material or ''
        if do_synthesize:
            for sec in ("技术前沿", "产业动态", "政策法规", "应用实例"):
                items = rendered.get(sec, [])
                for idx, item in enumerate(items):
                    user_prompt = data.get('prompt') or data.get('promptTitle') or None
                    user_prompt_title = data.get('promptTitle') if data.get('promptTitle') else None
                    snippet = _synthesize_fragment_by_ai(item.get('title', ''), cm_text, model=model, user_prompt=user_prompt, user_prompt_title=user_prompt_title)
                    if snippet:
                        try:
                            rendered[sec][idx]['summary'] = snippet
                        except Exception:
                            pass
    except Exception:
        logger.exception('synthesize snippets failed')

    # If client requested download, attempt advanced template rendering similar to original monolith
    if data.get('download'):
        now = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        out = os.path.join(tempfile.gettempdir(), f'report_{now}.docx')
        try:
            rendered_for_file = rendered
            for sec in ("技术前沿", "产业动态", "政策法规", "应用实例"):
                if not rendered_for_file.get(sec):
                    rendered_for_file[sec] = [{'title': '(暂无可用条目)', 'summary': '（该条暂无详细摘要 — 请检查 AI 输出或后端日志）'}]

            tpl_id = data.get('template_id') or data.get('templateId') or data.get('tpl_id')
            tpl_path = None
            if tpl_id is not None:
                try:
                    conn = get_db_conn()
                    c = conn.cursor()
                    c.execute('SELECT id, COALESCE(title, name) as title, content FROM prompt_template WHERE id=?', (int(tpl_id),))
                    row = c.fetchone()
                    conn.close()
                    if row:
                        tpl_path = row['content'] if 'content' in row.keys() else row.get('content')
                except Exception:
                    tpl_path = None

            used_docx_template = False
            if tpl_path and isinstance(tpl_path, str) and tpl_path.lower().endswith('.docx') and render_report:
                try:
                    if not os.path.isabs(tpl_path):
                        tpl_path = os.path.join(BASE_DIR, tpl_path)
                    if os.path.exists(tpl_path):
                        context = {
                            'title': data.get('title') or 'AI 生成报告',
                            'issue': data.get('issue') or '',
                            'org': data.get('org') or '',
                            'date': datetime.date.today().strftime('%Y-%m-%d'),
                            'core_titles': '\n'.join(rendered_for_file.get('core_news', [])),
                            'content': rendered_for_file
                        }
                        render_report(tpl_path, out, context)
                        used_docx_template = True
                except Exception:
                    used_docx_template = False

            if not used_docx_template and tpl_path and isinstance(tpl_path, str) and not tpl_path.lower().endswith('.docx'):
                try:
                    from docxtpl import DocxTemplate
                    from docx import Document as _Docx
                    base_tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
                    base_tmp.close()
                    doc = _Docx()
                    for line in tpl_path.splitlines() or ['']:
                        p = doc.add_paragraph()
                        p.add_run(line)
                    doc.save(base_tmp.name)
                    dt = DocxTemplate(base_tmp.name)
                    context = {
                        'title': data.get('title') or 'AI 生成报告',
                        'issue': data.get('issue') or '',
                        'org': data.get('org') or '',
                        'date': datetime.date.today().strftime('%Y-%m-%d'),
                        'core_titles': '\n'.join(rendered_for_file.get('core_news', [])),
                        'content': rendered_for_file
                    }
                    dt.render(context)
                    dt.save(out)
                    used_docx_template = True
                except Exception:
                    used_docx_template = False
                finally:
                    try:
                        if 'base_tmp' in locals() and os.path.exists(base_tmp.name):
                            os.remove(base_tmp.name)
                    except Exception:
                        pass

            if not used_docx_template:
                if generate_ai_report:
                    generate_ai_report(rendered_for_file, out, combined_material=combined_material)
                else:
                    # fall back
                    if simple_generate_docx:
                        simple_generate_docx(rendered_for_file, out)
                    else:
                        # fallback to _dump_docx_and_respond pathway
                        with open(out, 'w', encoding='utf-8') as f:
                            f.write(json.dumps(rendered_for_file or {}, ensure_ascii=False, indent=2))

            resp = send_file(out, as_attachment=True, download_name=os.path.basename(out))
            try:
                import json as _json
                resp.headers['X-AI-COUNTS'] = _json.dumps({k: len(v) for k, v in (rendered_for_file or {}).items()}, ensure_ascii=True)
            except Exception:
                pass
            return resp
        except Exception as e:
            logger.exception('生成报告失败')
            try:
                logger.debug('ai_content (partial): %s', str(ai_content)[:800])
            except Exception:
                pass
            return jsonify({'status': 'error', 'error': str(e)}), 500

    return jsonify({'status': 'ok', 'ai_content': ai_content, 'ai_counts': ai_counts})


def handle_generate_report_debug(req):
    # debug wrapper that echoes more request data
    resp = handle_generate_report(req)
    return resp


@generate_report_bp.route('/api/generate-report', methods=['POST'])
@require_token
def route_generate_report():
    return handle_generate_report(request)


@generate_report_bp.route('/api/generate-report-debug', methods=['GET', 'POST'])
@require_token
def route_generate_report_debug():
    return handle_generate_report_debug(request)
