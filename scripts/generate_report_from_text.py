#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 DOCX 报告：从纯文本中清理噪声（去除 deepseek 错误字串、控制字符、重复空行）并将内容写入 DOCX。
用法：
  python scripts/generate_report_from_text.py input.txt output/report_manual_generated.docx

依赖：python-docx
pip install python-docx
"""
import re
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

DEEPSEEK_ERROR_RE = re.compile(r"\{\'error\':\s*'call_deepseek_request_failed'.*?\}", re.S)
CONTROL_CHAR_RE = re.compile(r"[\x00-\x1f\x7f-\x9f]")
MULTI_EMPTY_RE = re.compile(r"\n{3,}")
FULLWIDTH_DIGIT_RE = re.compile(r"[０-９]")

CH_FONT = 'SimSun'  # 宋体（常见）
EN_FONT = 'Times New Roman'


def clean_text(text: str) -> str:
    # 移除 deepseek 错误字串
    text = DEEPSEEK_ERROR_RE.sub('', text)
    # 移除控制字符
    text = CONTROL_CHAR_RE.sub('', text)
    # 将全角数字转半角
    def fw2hw(m):
        return chr(ord(m.group(0)) - 0xFF10 + ord('0'))
    text = FULLWIDTH_DIGIT_RE.sub(fw2hw, text)
    # 规范化连续空行
    text = MULTI_EMPTY_RE.sub('\n\n', text)
    # 去掉行首尾空格
    text = '\n'.join([ln.strip() for ln in text.splitlines() if ln.strip() != ''])
    return text.strip()


def write_docx_from_text(text: str, outpath: str):
    doc = Document()

    style = doc.styles['Normal']
    try:
        style.font.name = EN_FONT
        style._element.rPr.rFonts.set(qn('w:eastAsia'), CH_FONT)
        style.font.size = Pt(11)
    except Exception:
        pass

    lines = text.split('\n')

    # 简单规则：以标题结尾符或冒号分割成章节
    sections = []
    cur_title = None
    cur_buf = []
    for ln in lines:
        if ln.endswith('：') or ln.endswith(':') or (len(ln) < 50 and ln.endswith('新闻智能体')):
            # treat as title
            if cur_title or cur_buf:
                sections.append((cur_title or '', '\n'.join(cur_buf).strip()))
            cur_title = ln.strip()
            cur_buf = []
        elif re.match(r'^第\d+期', ln) or re.match(r'^本期核心要闻', ln) or re.match(r'^[0-9]+\.', ln):
            # title-like
            if cur_title or cur_buf:
                sections.append((cur_title or '', '\n'.join(cur_buf).strip()))
            cur_title = ln.strip()
            cur_buf = []
        else:
            cur_buf.append(ln)
    if cur_title or cur_buf:
        sections.append((cur_title or '', '\n'.join(cur_buf).strip()))

    # 写入文档
    for title, body in sections:
        if title:
            p = doc.add_paragraph()
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = EN_FONT
            run._element.rPr.rFonts.set(qn('w:eastAsia'), CH_FONT)
        if body:
            for para in body.split('\n'):
                p = doc.add_paragraph()
                r = p.add_run(para)
                r.font.size = Pt(11)
                r.font.name = EN_FONT
                r._element.rPr.rFonts.set(qn('w:eastAsia'), CH_FONT)
        else:
            p = doc.add_paragraph()
            r = p.add_run('(无内容)')
            r.italic = True
            r._element.rPr.rFonts.set(qn('w:eastAsia'), CH_FONT)

    doc.save(outpath)


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Usage: python scripts/generate_report_from_text.py input.txt output.docx')
        sys.exit(2)
    with open(sys.argv[1], 'r', encoding='utf-8') as f:
        txt = f.read()
    cleaned = clean_text(txt)
    write_docx_from_text(cleaned, sys.argv[2])
    print('Saved', sys.argv[2])
